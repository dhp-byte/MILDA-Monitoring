################################################################################
# TABLEAU DE BORD AVANCÉ - Monitorage externe MILDA
# Version Premium avec Architecture Modulaire et Fonctionnalités Avancées
################################################################################

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx.shared import Inches
import scipy
from streamlit_folium import st_folium
import folium
import requests
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import simplekml
import random
from datetime import datetime, timedelta
import io
import os
import zipfile
import re
import math
import json
from io import BytesIO 
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import warnings
warnings.filterwarnings('ignore')

# Bibliothèques avancées
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from scipy import stats
    from sklearn.preprocessing import StandardScaler
    STATS_AVAILABLE = True
except ImportError:
    STATS_AVAILABLE = False

################################################################################
# CONFIGURATION GLOBALE
################################################################################

# Configuration de la page
st.set_page_config(
    page_title="MILDA Dashboard",
    page_icon="🦟",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://www.example.com/help',
        'Report a bug': "https://www.example.com/bug",
        'About': "# MILDA Dashboard v1.0\nTableau de bord pour le monitorage de la distribution des moustiquaires au Tchad "
    }
)

# Thème et styles personnalisés
CUSTOM_CSS = """
<style>
    /* En-tête principal */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    /* Cartes KPI */
    .kpi-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
        margin-bottom: 1rem;
    }
    
    .kpi-value {
        font-size: 2.5rem;
        font-weight: bold;
        color: #667eea;
        margin: 0;
    }
    
    .kpi-label {
        font-size: 0.9rem;
        color: #666;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .kpi-trend {
        font-size: 0.8rem;
        margin-top: 0.5rem;
    }
    
    .trend-up { color: #10b981; }
    .trend-down { color: #ef4444; }
    .trend-neutral { color: #6b7280; }
    
    /* Alertes */
    .alert-box {
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    .alert-success { background-color: #d1fae5; border-left: 4px solid #10b981; }
    .alert-warning { background-color: #fef3c7; border-left: 4px solid #f59e0b; }
    .alert-danger { background-color: #fee2e2; border-left: 4px solid #ef4444; }
    .alert-info { background-color: #dbeafe; border-left: 4px solid #3b82f6; }
    
    /* Filtres */
    .filter-section {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1.5rem;
    }
    
    /* Tables améliorées */
    .dataframe {
        font-size: 0.9rem !important;
    }
    
    /* Badges */
    .badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 12px;
        font-size: 0.8rem;
        font-weight: 600;
    }
    
    .badge-success { background-color: #d1fae5; color: #065f46; }
    .badge-warning { background-color: #fef3c7; color: #92400e; }
    .badge-danger { background-color: #fee2e2; color: #991b1b; }
    
    /* Animation de chargement */
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
    }
    
    .loading-pulse {
        animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
    }
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

################################################################################
# CLASSES ET STRUCTURES DE DONNÉES
################################################################################

class DataProcessor:
    """Classe pour le traitement avancé des données"""
    
    @staticmethod
    def normalize_yes_no(value):
        # Sécurité : si on reçoit un objet complexe (Series/List), on prend le premier élément ou on convertit
        if isinstance(value, (list, np.ndarray)):
            value = value[0] if len(value) > 0 else np.nan
            
        if pd.isna(value) or value == "":
            return "Non"
            #return 0
        
        val_str = str(value).lower().strip()
        # Votre nouveau formulaire utilise 'yes'/'no' en interne
        if val_str in ['oui', 'yes', '1', 'true']:
            return "Oui"
            #return 1
        return "Non"
        #return 0

    @staticmethod
    def normalize_norme(value):
        # Sécurité : si on reçoit un objet complexe (Series/List), on prend le premier élément ou on convertit
        if isinstance(value, (list, np.ndarray)):
            value = value[0] if len(value) > 0 else np.nan
            
        if pd.isna(value) or value == "":
            return "Non"
            #return 0
        
        #val_str = str(value).lower().strip()
        # Votre nouveau formulaire utilise 'yes'/'no' en interne
        if value in ['Oui – clé de répartition respectée', 'yes', '1', 'true']:
            return "Oui"
            #return 1
        return "Non"
        #return 0
    
    @staticmethod
    def calculate_expected_milda(nb_personnes: float) -> int:
        """Calcule le nombre requis selon votre nouvelle table"""
        if nb_personnes <= 2: return 1
        elif nb_personnes <= 4: return 2
        elif nb_personnes >= 5: return 3  # 5-6 pers = 3, et 7+ pers = 3
        return 0
    
    @staticmethod
    def clean_column_names(df: pd.DataFrame) -> pd.DataFrame:
        """Nettoie et normalise les noms de colonnes"""
        df = df.copy()
        df.columns = df.columns.str.strip().str.lower()
        df.columns = df.columns.str.replace(r'[^\w\s]', '_', regex=True)
        df.columns = df.columns.str.replace(r'\s+', '_', regex=True)
        return df
    
    @staticmethod
    def detect_outliers(series: pd.Series, method='iqr', threshold=1.5) -> pd.Series:
        """Détecte les valeurs aberrantes"""
        if method == 'iqr':
            Q1 = series.quantile(0.25)
            Q3 = series.quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - threshold * IQR
            upper_bound = Q3 + threshold * IQR
            return (series < lower_bound) | (series > upper_bound)
        elif method == 'zscore' and STATS_AVAILABLE:
            z_scores = np.abs(stats.zscore(series.dropna()))
            return z_scores > threshold
        return pd.Series([False] * len(series))


class MetricsCalculator:
    """Classe pour calculer les métriques et indicateurs"""
    
    @staticmethod
    def calculate_coverage_metrics(df: pd.DataFrame) -> Dict:
        """Calcule les métriques de couverture"""
        total_households = len(df)
        served = (df['indic_servi'] == 1).sum()
        correctly_served = (df['indic_correct'] == 1).sum()
        marked = (df['indic_marque'] == 1).sum()
        informed = (df['indic_info'] == 1).sum()
        
        return {
            'total_menages': total_households,
            'menages_servis': served,
            'menages_correct': correctly_served,
            'menages_marques': marked,
            'menages_informes': informed,
            'pct_servis': round(100 * served / total_households, 2) if total_households > 0 else 0,
            'pct_correct': round(100 * correctly_served / served, 2) if served > 0 else 0,
            'pct_marques': round(100 * marked / served, 2) if served > 0 else 0,
            'pct_informes': round(100 * informed / total_households, 2) if total_households > 0 else 0
        }
    
    @staticmethod
    def calculate_distribution_accuracy(df: pd.DataFrame) -> Dict:
        """Calcule la précision de la distribution"""
        df_served = df[df['menage_servi'] == 'Oui'].copy()
        
        if len(df_served) == 0:
            return {'precision': 0, 'sur_distribution': 0, 'sous_distribution': 0}
        
        df_served['ecart'] = df_served['nb_milda_recues'] - df_served['nb_milda_attendues']
        
        correct = (df_served['ecart'] == 0).sum()
        over = (df_served['ecart'] > 0).sum()
        under = (df_served['ecart'] < 0).sum()
        
        return {
            'precision': round(100 * correct / len(df_served), 2),
            'sur_distribution': round(100 * over / len(df_served), 2),
            'sous_distribution': round(100 * under / len(df_served), 2),
            'ecart_moyen': round(df_served['ecart'].mean(), 2)
        }
    
    @staticmethod
    def calculate_quality_score(metrics: Dict) -> float:
        """Calcule un score de qualité global (0-100)"""
        weights = {
            'pct_servis': 0.25,
            'pct_correct': 0.30,
            'pct_marques': 0.20,
            'pct_informes': 0.25
        }
        
        score = sum(metrics.get(k, 0) * w for k, w in weights.items())
        return round(score, 2)


class VisualizationEngine:
    """Classe pour créer des visualisations avancées"""
    
    COLOR_PALETTE = {
        'primary': '#667eea',
        'secondary': '#764ba2',
        'success': '#10b981',
        'warning': '#f59e0b',
        'danger': '#ef4444',
        'info': '#3b82f6'
    }
    
    @classmethod
    def create_kpi_gauge(cls, value: float, title: str, max_value: float = 100,
                        threshold_good: float = 80, threshold_medium: float = 60) -> go.Figure:
        """Crée un gauge KPI interactif"""
        
        # Déterminer la couleur selon les seuils
        if value >= threshold_good:
            color = cls.COLOR_PALETTE['success']
        elif value >= threshold_medium:
            color = cls.COLOR_PALETTE['warning']
        else:
            color = cls.COLOR_PALETTE['danger']
        
        fig = go.Figure(go.Indicator(
            mode="gauge+number+delta",
            value=value,
            title={'text': title, 'font': {'size': 16}},
            delta={'reference': threshold_good, 'increasing': {'color': cls.COLOR_PALETTE['success']}},
            gauge={
                'axis': {'range': [None, max_value], 'tickwidth': 1},
                'bar': {'color': color},
                'bgcolor': "white",
                'borderwidth': 2,
                'bordercolor': "gray",
                'steps': [
                    {'range': [0, threshold_medium], 'color': '#fee2e2'},
                    {'range': [threshold_medium, threshold_good], 'color': '#fef3c7'},
                    {'range': [threshold_good, max_value], 'color': '#d1fae5'}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': threshold_good
                }
            }
        ))
        
        fig.update_layout(
            height=250,
            margin=dict(l=20, r=20, t=50, b=20),
            paper_bgcolor="white",
            font={'color': "#333", 'family': "Arial"}
        )
        
        return fig
    
    @classmethod
    def create_stacked_bar_chart(cls, df: pd.DataFrame, x_col: str, 
                                 y_cols: List[str], title: str) -> go.Figure:
        """Crée un graphique à barres empilées avec annotations"""
        
        fig = go.Figure()
        
        colors = [cls.COLOR_PALETTE['primary'], cls.COLOR_PALETTE['success'], 
                 cls.COLOR_PALETTE['warning'], cls.COLOR_PALETTE['info']]
        
        for idx, col in enumerate(y_cols):
            fig.add_trace(go.Bar(
                name=col,
                x=df[x_col],
                y=df[col],
                marker_color=colors[idx % len(colors)],
                text=df[col].apply(lambda x: f'{x:.1f}%' if pd.notna(x) else ''),
                textposition='inside',
                textfont=dict(color='white', size=12),
                hovertemplate='<b>%{x}</b><br>' + col + ': %{y:.1f}%<extra></extra>'
            ))
        
        fig.update_layout(
            title={
                'text': f'<b>{title}</b>',
                'x': 0.5,
                'xanchor': 'center',
                'font': {'size': 20}
            },
            barmode='stack',
            xaxis_title="",
            yaxis_title="Pourcentage (%)",
            height=500,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="center",
                x=0.5
            ),
            template="plotly_white",
            hovermode='x unified'
        )
        
        return fig
    
    @classmethod
    def create_comparison_chart(cls, df: pd.DataFrame, categories: str, 
                               metric: str, title: str) -> go.Figure:
        """Crée un graphique de comparaison horizontal avec gradient"""
        
        df_sorted = df.sort_values(metric)
        
        # Créer un gradient de couleurs
        colors = [cls.COLOR_PALETTE['danger'] if v < 60 else 
                 cls.COLOR_PALETTE['warning'] if v < 80 else 
                 cls.COLOR_PALETTE['success'] for v in df_sorted[metric]]
        
        fig = go.Figure(go.Bar(
            y=df_sorted[categories],
            x=df_sorted[metric],
            orientation='h',
            marker=dict(
                color=colors,
                line=dict(color='rgba(0,0,0,0.3)', width=1)
            ),
            text=df_sorted[metric].apply(lambda x: f'{x:.1f}%'),
            textposition='outside',
            textfont=dict(size=12),
            hovertemplate='<b>%{y}</b><br>Valeur: %{x:.1f}%<extra></extra>'
        ))
        
        fig.update_layout(
            title={'text': f'<b>{title}</b>', 'x': 0.5, 'xanchor': 'center'},
            xaxis=dict(title="Pourcentage (%)", range=[0, 105]),
            yaxis=dict(title=""),
            height=max(400, len(df) * 40),
            template="plotly_white",
            showlegend=False
        )
        
        # Ajouter une ligne de référence à 80%
        fig.add_vline(x=80, line_dash="dash", line_color="red", 
                     annotation_text="Objectif 80%", annotation_position="top")
        
        return fig
    
    @classmethod
    def create_trend_chart(cls, df: pd.DataFrame, x_col: str, y_cols: List[str], 
                          title: str) -> go.Figure:
        """Crée un graphique de tendance avec lignes et marqueurs"""
        
        fig = go.Figure()
        
        colors = [cls.COLOR_PALETTE['primary'], cls.COLOR_PALETTE['success'], 
                 cls.COLOR_PALETTE['warning'], cls.COLOR_PALETTE['danger']]
        
        for idx, col in enumerate(y_cols):
            fig.add_trace(go.Scatter(
                x=df[x_col],
                y=df[col],
                mode='lines+markers',
                name=col,
                line=dict(color=colors[idx % len(colors)], width=3),
                marker=dict(size=8, symbol='circle'),
                hovertemplate='<b>%{x}</b><br>' + col + ': %{y:.1f}%<extra></extra>'
            ))
        
        fig.update_layout(
            title={'text': f'<b>{title}</b>', 'x': 0.5, 'xanchor': 'center'},
            xaxis_title="",
            yaxis_title="Pourcentage (%)",
            height=450,
            template="plotly_white",
            hovermode='x unified',
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5)
        )
        
        return fig
    
    @classmethod
    def create_heatmap(cls, df: pd.DataFrame, x_col: str, y_col: str, 
                      value_col: str, title: str) -> go.Figure:
        """Crée une heatmap interactive"""
        
        pivot_df = df.pivot_table(values=value_col, index=y_col, columns=x_col, aggfunc='mean')
        
        fig = go.Figure(data=go.Heatmap(
            z=pivot_df.values,
            x=pivot_df.columns,
            y=pivot_df.index,
            colorscale='RdYlGn',
            text=np.round(pivot_df.values, 1),
            texttemplate='%{text}%',
            textfont={"size": 10},
            hovertemplate='<b>%{y}</b><br>%{x}<br>Valeur: %{z:.1f}%<extra></extra>',
            colorbar=dict(title="Pourcentage (%)")
        ))
        
        fig.update_layout(
            title={'text': f'<b>{title}</b>', 'x': 0.5, 'xanchor': 'center'},
            height=max(400, len(pivot_df) * 30),
            template="plotly_white"
        )
        
        return fig
    
    @classmethod
    def create_sunburst_chart(cls, df: pd.DataFrame, path: List[str], 
                             values: str, title: str) -> go.Figure:
        """Crée un graphique sunburst hiérarchique"""
        
        fig = px.sunburst(
            df,
            path=path,
            values=values,
            color=values,
            color_continuous_scale='RdYlGn',
            title=f'<b>{title}</b>'
        )
        
        fig.update_layout(
            height=600,
            template="plotly_white"
        )
        
        return fig


class ReportGenerator:
    """Classe pour générer des rapports dans différents formats"""
    
    @staticmethod
    def generate_excel_report(data: pd.DataFrame, tables: Dict[str, pd.DataFrame], 
                             metrics: Dict) -> io.BytesIO:
        """Génère un rapport Excel multi-feuilles"""
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Feuille de résumé
            summary_df = pd.DataFrame([metrics])
            summary_df.to_excel(writer, sheet_name='Résumé', index=False)
            
            # Données brutes
            data.to_excel(writer, sheet_name='Données brutes', index=False)
            
            # Tableaux d'analyse
            for name, table in tables.items():
                sheet_name = name.replace('_', ' ').title()[:31]  # Excel limite à 31 caractères
                table.to_excel(writer, sheet_name=sheet_name, index=False)
        
        output.seek(0)
        return output
    
    @staticmethod
    def generate_json_report(data: pd.DataFrame, metrics: Dict) -> str:
        """Génère un rapport JSON"""
        report = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'total_records': len(data),
                'version': '2.0'
            },
            'metrics': metrics,
            'data_summary': {
                'provinces': data['province'].unique().tolist() if 'province' in data.columns else [],
                'districts': data['district'].unique().tolist() if 'district' in data.columns else []
            }
        }
        
        return json.dumps(report, ensure_ascii=False, indent=2)

def load_github_mappings(url):
    """
    Télécharge le fichier Excel depuis GitHub et crée un dictionnaire 
    de mapping structuré par list_name.
    """
    try:
        # 1. Récupération du fichier avec un timeout pour éviter de bloquer l'app
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        
        # 2. Lecture du flux binaire Excel
        # On force toutes les colonnes en 'str' pour éviter que les codes 
        # numériques ne soient interprétés comme des nombres ou des floats.
        df_choices = pd.read_excel(
            BytesIO(response.content), 
            sheet_name='Choix', 
            dtype=str
        )
        
        # 3. Nettoyage des en-têtes de colonnes
        df_choices.columns = df_choices.columns.str.strip()
        
        # 4. Suppression des lignes totalement vides ou incomplètes
        # (Indispensable si le fichier Excel a des lignes fantômes en bas)
        df_choices = df_choices.dropna(subset=['list_name', 'value'])
        
        # 5. Construction du dictionnaire de dictionnaires
        # Structure : { 'nom_liste': { 'code': 'label' } }
        mappings = {}
        
        # On boucle sur chaque catégorie unique (province, district, sexe, etc.)
        for list_name in df_choices['list_name'].unique():
            # Nettoyage du nom de la liste
            clean_list_key = str(list_name).strip()
            
            # Extraction du sous-ensemble correspondant à cette liste
            subset = df_choices[df_choices['list_name'] == list_name]
            
            # Création du mapping interne (Code -> Libellé)
            # On applique .strip() sur les valeurs et labels pour éviter "1 " != "1"
            mappings[clean_list_key] = dict(zip(
                subset['value'].str.strip(), 
                subset['label'].str.strip()
            ))
        
        # Petit message de succès pour le debug (visible uniquement en console)
        print(f"✅ Mapping chargé avec succès : {len(mappings)} listes extraites.")
        
        return mappings

    except requests.exceptions.RequestException as e:
        st.error(f"❌ Erreur de connexion au dépôt GitHub : {e}")
        return None
    except ValueError as e:
        st.error(f"❌ Erreur de format : Vérifiez que l'onglet s'appelle 'Choices'. Détails : {e}")
        return None
    except Exception as e:
        st.error(f"❌ Une erreur inattendue est survenue : {e}")
        return None


def load_multiple_github_mappings(selected_phases: List[str]) -> Dict:
    """
    Télécharge et fusionne les dictionnaires de mapping pour toutes les phases sélectionnées.
    """
    combined_mappings = {}
    
    if not selected_phases:
        return combined_mappings

    for phase in selected_phases:
        url = MAPPINGS_URLS.get(phase)
        if not url:
            continue
            
        try:
            response = requests.get(url, timeout=15)
            response.raise_for_status()
            
            df_choices = pd.read_excel(BytesIO(response.content), sheet_name='Choix', dtype=str)
            df_choices.columns = df_choices.columns.str.strip()
            df_choices = df_choices.dropna(subset=['list_name', 'value'])
            
            for list_name in df_choices['list_name'].unique():
                clean_list_key = str(list_name).strip()
                subset = df_choices[df_choices['list_name'] == list_name]
                
                # Nouveau dictionnaire pour cette phase
                phase_dict = dict(zip(subset['value'].str.strip(), subset['label'].str.strip()))
                
                # Si la catégorie existe déjà (ex: 'province'), on fusionne les dictionnaires
                if clean_list_key in combined_mappings:
                    combined_mappings[clean_list_key].update(phase_dict)
                else:
                    combined_mappings[clean_list_key] = phase_dict
                    
        except Exception as e:
            st.error(f"❌ Erreur lors du chargement des mappings pour la {phase} : {e}")
            
    return combined_mappings       
# URL vers votre fichier (format RAW)
# URLs vers les fichiers d'onglets 'Choix' de chaque phase (Format RAW requis)
MAPPINGS_URLS = {
    "Phase 1": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_1.xlsx",
    "Phase 2": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_2.xlsx",
    "Phase 3": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_3.xlsx" # À adapter si le nom change
}
mappings = load_multiple_github_mappings(MAPPINGS_URLS)

################################################################################
# FONCTIONS DE TRAITEMENT DES DONNÉES
################################################################################
def process_milda_dataframe(data: pd.DataFrame, mappings_dict: Dict = None) -> Tuple[pd.DataFrame, Dict]:
    """Applique la logique métier commune à Excel et KoBo.
    mappings_dict : dictionnaire de correspondances codes->labels pour la phase concernée.
    Si None, utilise st.session_state.mappings ou la variable globale mappings.
    """
    if mappings_dict is None:
        mappings_dict = st.session_state.get('mappings') or mappings or {}
    # 1. Mapping des colonnes (votre dictionnaire existant)
    column_mapping = {
            'province': ['province', 'Province', 'S0Q02'],
            'district': ['district', 'district sanitaire', 'District sanitaire de :', 'S0Q06'],
            'centre_sante': ['centre_sante', 'centre de santé', 'Centre de santé', 'S0Q07'],
            'date_enquete': ['date_enquete', 'date_enquête', 'Date enquête', 'Date', 'Date de l’enquête', 'S0Q01'],
            'start': ['start'],
            'sexe': ['S1Q14', 'Sexe du répondant', 'Sexe', 'sexe'],
            'consentement': ['S1Q01'],
            'activ_rev': ['S1Q05', 'Profession du chef de ménage'],
            'heure_interview': ['heure_interview', 'Heure', 'time', 'heure', 'end'], 
            'agent_name': ['agent_name', "Nom de l'enquêteur", 'Enquêteur', 'Username', 'S0Q05'],
            #'enqueteur': ['enqueteur'],
            'village': ['village', 'Village/Avenue/Quartier', 'S0Q08'],
            'menage_chef' : ['S1Q02', 'Etes-vous le Chef de ce ménage ?', 'gr_1/S1Q2'],
            'menage_servi': ['Est-ce que le ménage a-t-il été servi en MILDA lors de la campagne de distribution de masse ?', 'gr_1/S1Q17', 'S1Q17' ],
            'nb_personnes': ['nb_personnes', 'Nombre des personnes qui habitent dans le ménage', 'gr_1/S1Q19', 'S1Q19'],
            'nb_milda_recues': ['nb_milda_recues', 'Combien de MILDA avez-vous reçues ?', 'gr_1/S1Q20', 'S1Q20'],
            #'verif_cle': ['verif_cle', 'gr_1/verif_cle', 'verif_cle'],
            #'norme': ['norme', 'gr_1/S1Q21', 'S1Q21'],
            'norme': ['verif_cle', 'gr_1/verif_cle', 'verif_cle'],
            'menage_marque': ['menage_marque', 'Est-ce que le ménage a  été marqué comme un ménage ayant reçu de MILDA?', 'gr_1/S1Q22', 'S1Q22'],
            'sensibilise': ['sensibilise', 'Avez-vous été sensibilisé sur l’utilisation correcte du MILDA par les relais communautaires ?', 'gr_1/S1Q23', 'S1Q23'],
            'latitude': ['latitude', '_LES COORDONNEES GEOGRAPHIQUES_latitude', '_geolocation'],
            'longitude': ['longitude', '_LES COORDONNEES GEOGRAPHIQUES_longitude'],
            'respondant_col' : ['S1Q18', 'Le répondant est-il le même que lors de la distribution ?'],
            'id_scan' : ['scan_milda', 'Scannage code QR MILDA', '${agent_name}, Avez pas pu scanner un nombre codes QR corresondant aux MILDA reçu dans le ménage?', 'rsn2'],
            'raison' : ['Sélectionner la raison', 'S1Q25'],
            'raison_scan' : ["${agent_name},Pourquoi vous n'avez pas pu scanner nombre codes QR corresondant aux MILDA reçu dans le ménage?", 'rsn'],
            'source': ['Où avez-vous vu ou entendu ces informations ?', 'source'],
            'conseil' : ['sensibilisation', "Au cours du mois dernier, quelles instructions d'utilisation et d'entretien des moustiquaires avez-vous vues ou entendues?"],
            'information' : ['Étiez-vous informé qu’il y aurait une campagne de distribution de moustiquaires et que des agents visiteraient les ménages ?', 'information']
        }

    # Nettoyage des noms (enlève les préfixes gr_1/ etc.)
    data.columns = [c.split('/')[-1] for c in data.columns]
    
    # Application du mapping
    rename_dict = {}
    for target, sources in column_mapping.items():
        for source in sources:
            if source in data.columns:
                rename_dict[source] = target
                break
    data = data.rename(columns=rename_dict)

    if mappings_dict:
        # Configuration étendue avec vos nouvelles variables
        config = {
            'province': 'province',
            'district': 'district',
            'centre_sante': 'cs',
            'village': 'village',
            #'agent_name' : 'enqueteur',
            'sexe': 'sexe',
            'activ_rev': 'activ_rev',
            'raison': 'raison',          # Choix multiples possibles
            'conseil': 'conseil',        # Choix multiples possibles
            'source': 'source',          # Choix multiples possibles
            'raison_scan': 'raison_scan'
        }

        # Liste spécifique des colonnes qui peuvent avoir plusieurs réponses (Select Multiple)
        multi_choice_cols = ['raison', 'conseil', 'source']

        for col, list_name in config.items():
            if col in data.columns:
                # Nettoyage de base — CORRECTION : utiliser str.replace() et pas .replace()
                data[col] = (
                    data[col]
                    .astype(str)
                    .str.replace(r'\.0$', '', regex=True)
                    .str.strip()
                    .str.replace(r'^nan$', '', regex=True)  # ← CORRIGÉ : str.replace() au lieu de .replace()
                )
                
                if list_name in mappings_dict:
                    if col in multi_choice_cols:
                        # --- LOGIQUE POUR CHOIX MULTIPLES ---
                        def decode_multi(val, mapping_dict):
                            if not val or val == '': return val
                            codes = str(val).split()
                            labels = [mapping_dict.get(c, c) for c in codes]
                            return ", ".join(labels)
                        
                        data[col] = data[col].apply(lambda x: decode_multi(x, mappings_dict[list_name]))
                    else:
                        # --- LOGIQUE POUR CHOIX UNIQUE ---
                        data[col] = data[col].replace(mappings_dict[list_name])
                
    # Traitement spécial GPS pour KoBo (si format liste [lat, long])
    if 'latitude' in data.columns and isinstance(data['latitude'].iloc[0], list):
        coords = data['latitude']
        data['latitude'] = coords.apply(lambda x: x[0] if isinstance(x, list) else None)
        data['longitude'] = coords.apply(lambda x: x[1] if isinstance(x, list) else None)
        
    
    # Normalisation Oui/Non
    # Dans process_milda_dataframe
    cols_to_fix = ['menage_servi', 'menage_marque', 'information', 'menage_chef', 'respondant_col', 'sensibilise', 'consentement']
    for col in cols_to_fix:
        if col in data.columns:
            # On s'assure de ne traiter que des valeurs simples
            data[col] = data[col].astype(str).apply(DataProcessor.normalize_yes_no)

    data['norme'] = data['norme'].apply(DataProcessor.normalize_norme)
    data = data[data['consentement'] == 'Oui']
    # Conversions numériques et indicateurs
    for col in ['nb_personnes', 'nb_milda_recues']:
        if col in data.columns:
            data[col] = pd.to_numeric(data[col], errors='coerce').fillna(0)

    data['nb_milda_attendues'] = data['nb_personnes'].apply(DataProcessor.calculate_expected_milda)
    data['ecart_distribution'] = data['nb_milda_recues'] - data['nb_milda_attendues']
    
    # Indicateurs binaires pour le Dashboard
    # Remplacez les lignes 207-210 par :
    
    # APRÈS (corrigé)
    # Forcez la création des colonnes indicateurs même si les données sont brutes
    # Note : on compare avec 'yes' car c'est la valeur interne de votre nouveau formulaire
    #data['indic_servi'] = (data.get('menage_servi', pd.Series()).astype(str).str.lower() == 'yes').astype(int)
    #data['indic_marque'] = (data.get('menage_marque', pd.Series()).astype(str).str.lower() == 'yes').astype(int)
    
    # Calcul du "Correctement servi" (Servi + Norme respectée)
    #norme_ok = (data.get('norme', pd.Series()).astype(str).str.lower() == 'yes')
    #data['indic_correct'] = ((data['indic_servi'] == 1) & norme_ok).astype(int)
    data['indic_servi'] = (data['menage_servi'] == 'Oui').astype(int)
    #data['indic_correct'] = ((data['menage_servi'] == 'Oui') & (data.get('norme') == 'Oui')).astype(int)
    data['indic_correct'] = ((data['norme'] == 'Oui')).astype(int)
    data['indic_marque'] = (data['menage_marque'] == 'Oui').astype(int)
    data['indic_info'] = (data['information'] == 'Oui').astype(int)

    if 'date_enquete' in data.columns:
        data['date_enquete'] = pd.to_datetime(data['date_enquete'], errors='coerce')

    stats = {
        'total_rows': len(data),
        'total_provinces': data['province'].nunique() if 'province' in data.columns else 0,
        'date_range': (data['date_enquete'].min(), data['date_enquete'].max())
    }
    return data, stats

@st.cache_data(ttl=600, show_spinner=False) # Cache plus court pour KoBo (10 min)
def load_data_from_kobo(server_url: str, asset_uid: str, token: str) -> Tuple[pd.DataFrame, Dict]:
    """Récupère les données via l'API KoBo et les traite"""
    try:
        headers = {"Authorization": f"Token {token}"}
        url = f"{server_url}/api/v2/assets/{asset_uid}/data.json"
        
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            results = response.json().get('results', [])
            if not results:
                return pd.DataFrame(), {}
            
            df_raw = pd.DataFrame(results)
            # On appelle la fonction de traitement universelle
            return process_milda_dataframe(df_raw)
        else:
            st.error(f"Erreur API KoBo : {response.status_code}")
            return pd.DataFrame(), {}
            
    except Exception as e:
        st.error(f"Erreur de connexion KoBo : {str(e)}")
        return pd.DataFrame(), {}
        

@st.cache_data(ttl=3600, show_spinner=False)
def load_and_process_data(uploaded_file, sheet_name: str = None) -> Tuple[pd.DataFrame, Dict]:
    """Charge et traite les données Excel en alignant les indicateurs sur le format KoBo"""
    
    try:
        # 1. Lecture du fichier
        if sheet_name:
            try:
                data = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            except:
                data = pd.read_excel(uploaded_file, sheet_name=0)
        else:
            data = pd.read_excel(uploaded_file, sheet_name=0)
        
        # 2. Mapping des colonnes (Identique à KoBo pour l'unification)
        column_mapping = {
            'province': ['province', 'Province', 'S0Q02'],
            'district': ['district', 'district sanitaire', 'District sanitaire de :', 'S0Q06'],
            'centre_sante': ['centre_sante', 'centre de santé', 'Centre de santé', 'S0Q07', 'cs'],
            'date_enquete': ['date_enquete', 'date_enquête', 'Date enquête', 'Date', 'Date de l’enquête', 'S0Q01'],
            'agent_name': ['agent_name', "Nom de l'enquêteur", 'Enquêteur', 'Username', 'S0Q05', 'enqueteur'],
            'village': ['village', 'Village/Avenue/Quartier', 'S0Q08'],
            'sexe': ['S1Q14', 'Sexe du répondant', 'Sexe', 'sexe'],
            'activ_rev': ['S1Q05', 'Profession du chef de ménage', 'activ_rev'],
            'menage_chef' : ['S1Q02', 'Etes-vous le Chef de ce ménage ?', 'menage_chef'],
            'menage_servi': ['Est-ce que le ménage a-t-il été servi en MILDA lors de la campagne de distribution de masse ?', 'S1Q17', 'menage_servi'],
            'norme': ['norme', 'S1Q21', 'respect_norme', "${agent_name},Ce ménage a-t-il été servi conformément à la norme de la CDM 2026?"],
            'menage_marque': ['menage_marque', 'Est-ce que le ménage a  été marqué comme un ménage ayant reçu de MILDA?', 'S1Q22'],
            'information': ['Étiez-vous informé qu’il y aurait une campagne de distribution de moustiquaires et que des agents visiteraient les ménages ?', 'information', 'S1Q23'],
            'nb_personnes': ['nb_personnes', 'Nombre des personnes qui habitent dans le ménage', 'S1Q19'],
            'nb_milda_recues': ['nb_milda_recues', 'Combien de MILDA avez-vous reçues ?', 'S1Q20'],
            'latitude': ['_LES COORDONNEES GEOGRAPHIQUES_latitude', 'latitude','lat'],
            'longitude': ['_LES COORDONNEES GEOGRAPHIQUES_longitude', 'longitude', 'long'],
            'end_dt': ['end', 'Heure de fin', 'end_dt']
        }
        
        # Application du renommage
        rename_dict = {}
        for target, sources in column_mapping.items():
            for source in sources:
                if source in data.columns:
                    rename_dict[source] = target
                    break
        data = data.rename(columns=rename_dict)

    
        # 3. Nettoyage et Normalisation (Crucial pour le rapport)
        # On s'assure que 'Oui'/'Non' sont uniformes
        yes_no_cols = ['menage_servi', 'norme', 'menage_marque', 'information', 'menage_chef']
        for col in yes_no_cols:
            if col in data.columns:
                data[col] = data[col].astype(str).apply(DataProcessor.normalize_yes_no)
        
        # 4. Conversions numériques
        for col in ['nb_personnes', 'nb_milda_recues']:
            if col in data.columns:
                data[col] = pd.to_numeric(data[col], errors='coerce').fillna(0)

        # APRÈS avoir appliqué le rename_dict, ajoutez cette sécurité :
        if 'end_dt' not in data.columns:
            if 'date_enquete' in data.columns:
                data['end_dt'] = data['date_enquete'] # Repli sur la date d'enquête
            else:
                data['end_dt'] = pd.Timestamp.now()
        # 5. Calcul des indicateurs de performance (ALIGNÉS SUR LE RAPPORT WORD)
        # Ces colonnes 'indic_...' sont celles utilisées par generate_automatic_report
        data['indic_servi'] = (data['menage_servi'] == 'Oui').astype(int)
        
        # Correctement servi = Servi ET Norme respectée
        data['indic_correct'] = ((data['menage_servi'] == 'Oui') & (data.get('norme') == 'Oui')).astype(int)
        
        # Marqué = Ménage marqué 'Oui'
        data['indic_marque'] = (data['menage_marque'] == 'Oui').astype(int)
        
        # Information = Informé 'Oui'
        data['indic_info'] = (data.get('information') == 'Oui').astype(int)

        # 6. Calcul des besoins (Règle CDM 2026)
        if 'nb_personnes' in data.columns:
            data['nb_milda_attendues'] = data['nb_personnes'].apply(DataProcessor.calculate_expected_milda)
            if 'nb_milda_recues' in data.columns:
                data['ecart_distribution'] = data['nb_milda_recues'] - data['nb_milda_attendues']
        
        # 7. Traitement des dates pour le titre du rapport
        if 'date_enquete' in data.columns:
            data['date_enquete'] = pd.to_datetime(data['date_enquete'], errors='coerce')
        
        # 8. Statistiques pour l'interface Streamlit
        stats = {
            'total_rows': len(data),
            'total_provinces': data['province'].nunique() if 'province' in data.columns else 0,
            'total_districts': data['district'].nunique() if 'district' in data.columns else 0,
            'date_range': (
                data['date_enquete'].min().strftime('%d/%m/%Y') if 'date_enquete' in data.columns and data['date_enquete'].notna().any() else 'N/A',
                data['date_enquete'].max().strftime('%d/%m/%Y') if 'date_enquete' in data.columns and data['date_enquete'].notna().any() else 'N/A'
            )
        }
        
        return data, stats
        
    except Exception as e:
        st.error(f"Erreur lors du traitement Excel : {str(e)}")
        return pd.DataFrame(), {}


@st.cache_data(show_spinner=False)
def generate_analysis_tables(data: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """Génère les tableaux d'analyse"""
    
    tables = {}
    
    # Table 0: Résumé par province
    if 'province' in data.columns:
        tables['resume_province'] = data.groupby('province').agg(
            menages_total=('province', 'count'),
            menages_servis=('indic_servi', 'sum'),
            menages_correct=('indic_correct', 'sum'),
            menages_marques=('indic_marque', 'sum'),
            menages_informes=('indic_info', 'sum'),
            pct_servis=('indic_servi', lambda x: round(100 * x.mean(), 1)),
            pct_correct=('indic_correct', lambda x: round(100 * x.mean(), 1)),
            pct_marques=('indic_marque', lambda x: round(100 * x.mean(), 1)),
            pct_informes=('indic_info', lambda x: round(100 * x.mean(), 1))
        ).reset_index()
    
    # Table 1: Détail par district
    if 'province' in data.columns and 'district' in data.columns:
        tables['detail_district'] = data.groupby(['province', 'district']).agg(
            menages_total=('district', 'count'),
            pct_servis=('indic_servi', lambda x: round(100 * x.mean(), 1)),
            pct_correct=('indic_correct', lambda x: round(100 * x.mean(), 1)),
            pct_marques=('indic_marque', lambda x: round(100 * x.mean(), 1)),
            pct_informes=('indic_info', lambda x: round(100 * x.mean(), 1))
        ).reset_index()
    
    # Table 2: Analyse de la distribution
    if 'ecart_distribution' in data.columns:
        distribution_df = data[data['menage_servi'] == 'Oui'].copy()
        if len(distribution_df) > 0:
            tables['analyse_distribution'] = distribution_df.groupby('province').agg(
                total=('province', 'count'),
                distribution_exacte=('ecart_distribution', lambda x: round(100 * (x == 0).mean(), 1)),
                sur_distribution=('ecart_distribution', lambda x: round(100 * (x > 0).mean(), 1)),
                sous_distribution=('ecart_distribution', lambda x: round(100 * (x < 0).mean(), 1)),
                ecart_moyen=('ecart_distribution', lambda x: round(x.mean(), 2))
            ).reset_index()
    
    # Table 3: Performance par enquêteur
    if 'agent_name' in data.columns:
        tables['performance_enqueteur'] = data.groupby('agent_name').agg(
            nombre_enquetes=('agent_name', 'count'),
            pct_servis=('indic_servi', lambda x: round(100 * x.mean(), 1)),
            pct_correct=('indic_correct', lambda x: round(100 * x.mean(), 1)),
            qualite_score=('indic_correct', lambda x: round(100 * x.mean(), 1))
        ).reset_index().sort_values('qualite_score', ascending=False)
    
    return tables


################################################################################
# INTERFACE UTILISATEUR - PAGES
################################################################################

def render_header():
    """Affiche l'en-tête principal"""
    st.markdown("""
        <div class="main-header">
            <h1>🦟 MILDA Dashboard</h1>
            <p style="font-size: 1.2rem; margin-top: 0.5rem;">
                Système de monitorage et d'analyse de la distribution des moustiquaires au Tchad 2026
            </p>
        </div>
    """, unsafe_allow_html=True)


def render_kpi_cards(metrics: Dict):
    """Affiche les cartes KPI principales"""
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
            <div class="kpi-card">
                <p class="kpi-label">Ménages Servis</p>
                <p class="kpi-value">{metrics.get('pct_servis', 0):.1f}%</p>
                <p class="kpi-trend trend-{'up' if metrics.get('pct_servis', 0) >= 80 else 'down'}">
                    {'✓ Objectif atteint' if metrics.get('pct_servis', 0) >= 80 else '⚠ Sous objectif'}
                </p>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
            <div class="kpi-card">
                <p class="kpi-label">Distribution Correcte</p>
                <p class="kpi-value">{metrics.get('pct_correct', 0):.1f}%</p>
                <p class="kpi-trend trend-{'up' if metrics.get('pct_correct', 0) >= 80 else 'down'}">
                    {'✓ Objectif atteint' if metrics.get('pct_correct', 0) >= 80 else '⚠ Sous objectif'}
                </p>
            </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
            <div class="kpi-card">
                <p class="kpi-label">Ménages Marqués</p>
                <p class="kpi-value">{metrics.get('pct_marques', 0):.1f}%</p>
                <p class="kpi-trend trend-{'up' if metrics.get('pct_marques', 0) >= 80 else 'down'}">
                    {'✓ Objectif atteint' if metrics.get('pct_marques', 0) >= 80 else '⚠ Sous objectif'}
                </p>
            </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
            <div class="kpi-card">
                <p class="kpi-label">Ménages Informés</p>
                <p class="kpi-value">{metrics.get('pct_informes', 0):.1f}%</p>
                <p class="kpi-trend trend-{'up' if metrics.get('pct_informes', 0) >= 80 else 'down'}">
                    {'✓ Objectif atteint' if metrics.get('pct_informes', 0) >= 80 else '⚠ Sous objectif'}
                </p>
            </div>
        """, unsafe_allow_html=True)


def render_alerts(metrics: Dict):
    """Affiche les alertes basées sur les seuils"""
    
    alerts = []
    
    if metrics.get('pct_servis', 0) < 70:
        alerts.append(('danger', f"Taux de ménages servis critique: {metrics['pct_servis']:.1f}% (objectif: 80%)"))
    elif metrics.get('pct_servis', 0) < 80:
        alerts.append(('warning', f"Taux de ménages servis sous l'objectif: {metrics['pct_servis']:.1f}% (objectif: 80%)"))
    else:
        alerts.append(('success', f"Excellent taux de ménages servis: {metrics['pct_servis']:.1f}%"))
    
    if metrics.get('pct_correct', 0) < 70:
        alerts.append(('danger', f"Précision de distribution critique: {metrics['pct_correct']:.1f}%"))
    
    if metrics.get('pct_informes', 0) < 60:
        alerts.append(('warning', "Sensibilisation insuffisante sur l'utilisation des MILDA"))
    
    for alert_type, message in alerts:
        st.markdown(f"""
            <div class="alert-box alert-{alert_type}">
                <strong>{'🔴' if alert_type == 'danger' else '⚠️' if alert_type == 'warning' else '✅'}</strong> {message}
            </div>
        """, unsafe_allow_html=True)


def page_dashboard(data: pd.DataFrame, tables: Dict[str, pd.DataFrame]):
    """Page principale du dashboard"""
    
    st.markdown("## 📊 Vue d'ensemble")
    
    # Calcul des métriques
    metrics = MetricsCalculator.calculate_coverage_metrics(data)
    quality_score = MetricsCalculator.calculate_quality_score(metrics)
    
    # KPIs principaux
    render_kpi_cards(metrics)
    
    st.markdown("---")
    
    # Alertes
    with st.expander("🔔 Alertes et Notifications", expanded=True):
        render_alerts(metrics)
    
    st.markdown("---")
    
    # Graphiques principaux
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📈 Indicateurs par Province")
        if 'resume_province' in tables and len(tables['resume_province']) > 0:
            fig = VisualizationEngine.create_comparison_chart(
                tables['resume_province'],
                'province',
                'pct_servis',
                'Taux de couverture par province'
            )
            st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        st.markdown("### 🎯 Score de qualité global")
        fig = VisualizationEngine.create_kpi_gauge(
            quality_score,
            "Score de qualité",
            max_value=100,
            threshold_good=80,
            threshold_medium=60
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # Graphiques secondaires
    if 'resume_province' in tables:
        st.markdown("### 📊 Comparaison des indicateurs")
        fig = VisualizationEngine.create_stacked_bar_chart(
            tables['resume_province'],
            'province',
            ['pct_servis', 'pct_correct', 'pct_marques', 'pct_informes'],
            'Indicateurs de qualité par province'
        )
        st.plotly_chart(fig, use_container_width=True)


def page_analysis(data: pd.DataFrame, tables: Dict[str, pd.DataFrame]):
    """Page d'analyse détaillée"""
    
    st.markdown("## 🔍 Analyse Détaillée")
    
    # Filtres
    st.markdown('<div class="filter-section">', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    
    with col1:
        provinces = ['Toutes'] + sorted(data['province'].dropna().unique().tolist())
        selected_province = st.selectbox("🗺️ Province", provinces)
    
    with col2:
        if selected_province != 'Toutes':
            districts = ['Tous'] + sorted(data[data['province'] == selected_province]['district'].dropna().unique().tolist())
        else:
            districts = ['Tous'] + sorted(data['district'].dropna().unique().tolist())
        selected_district = st.selectbox("📍 District", districts)
    
    with col3:
        date_range = st.date_input(
            "📅 Période",
            value=(data['date_enquete'].min(), data['date_enquete'].max()) if 'date_enquete' in data.columns else (datetime.now(), datetime.now()),
            key='date_filter'
        )
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Filtrer les données
    filtered_data = data.copy()
    if selected_province != 'Toutes':
        filtered_data = filtered_data[filtered_data['province'] == selected_province]
    if selected_district != 'Tous':
        filtered_data = filtered_data[filtered_data['district'] == selected_district]
    
    # Métriques filtrées
    st.markdown("### 📈 Métriques de la sélection")
    filtered_metrics = MetricsCalculator.calculate_coverage_metrics(filtered_data)
    
    col1, col2, col3, col4, col5 = st.columns(5)
    col1.metric("Ménages analysés", filtered_metrics['total_menages'])
    col2.metric("Servis", f"{filtered_metrics['pct_servis']:.1f}%")
    col3.metric("Correct", f"{filtered_metrics['pct_correct']:.1f}%")
    col4.metric("Marqués", f"{filtered_metrics['pct_marques']:.1f}%")
    col5.metric("Informés", f"{filtered_metrics['pct_informes']:.1f}%")
    
    st.markdown("---")
    
    # Analyse de la distribution
    st.markdown("### 📦 Analyse de la distribution")
    dist_metrics = MetricsCalculator.calculate_distribution_accuracy(filtered_data)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"""
            <div class="kpi-card">
                <h4>Précision de distribution</h4>
                <p class="kpi-value">{dist_metrics['precision']:.1f}%</p>
                <p>Distribution exacte selon la norme</p>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div class="kpi-card">
                <h4>Écart moyen</h4>
                <p class="kpi-value">{dist_metrics['ecart_moyen']:.2f}</p>
                <p>MILDA par ménage (écart à la norme)</p>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        # Graphique de répartition des écarts
        ecart_data = pd.DataFrame({
            'Type': ['Distribution exacte', 'Sur-distribution', 'Sous-distribution'],
            'Pourcentage': [dist_metrics['precision'], dist_metrics['sur_distribution'], dist_metrics['sous_distribution']]
        })
        
        fig = px.pie(
            ecart_data,
            values='Pourcentage',
            names='Type',
            title='<b>Répartition des types de distribution</b>',
            color_discrete_sequence=[VisualizationEngine.COLOR_PALETTE['success'], 
                                    VisualizationEngine.COLOR_PALETTE['warning'], 
                                    VisualizationEngine.COLOR_PALETTE['danger']]
        )
        fig.update_layout(height=350)
        st.plotly_chart(fig, use_container_width=True)
    
    st.markdown("---")
    
    # Tableaux détaillés
    st.markdown("### 📋 Tableaux détaillés")
    
    tab1, tab2, tab3 = st.tabs(["Par District", "Par Enquêteur", "Distribution"])
    
    with tab1:
        if 'detail_district' in tables:
            display_table = tables['detail_district']
            if selected_province != 'Toutes':
                display_table = display_table[display_table['province'] == selected_province]
            
            st.dataframe(
                display_table.style.background_gradient(
                    subset=['pct_servis', 'pct_correct', 'pct_marques', 'pct_informes'],
                    cmap='RdYlGn',
                    vmin=0,
                    vmax=100
                ),
                use_container_width=True
            )

    
    with tab2:
        if 'performance_enqueteur' in tables:
            st.dataframe(
                tables['performance_enqueteur'].style.background_gradient(
                    subset=['qualite_score'],
                    cmap='RdYlGn',
                    vmin=0,
                    vmax=100
                ),
                use_container_width=True
            )
    
    with tab3:
        if 'analyse_distribution' in tables:
            st.dataframe(
                tables['analyse_distribution'].style.background_gradient(
                    subset=['distribution_exacte'],
                    cmap='RdYlGn',
                    vmin=0,
                    vmax=100
                ),
                use_container_width=True
            )


def page_maps(data: pd.DataFrame):
    """Page avec visualisations géographiques"""
    
    st.markdown("## 🗺️ Visualisation Géographique")
    
    if 'latitude' not in data.columns or 'longitude' not in data.columns:
        st.warning("Données de géolocalisation non disponibles dans le fichier")
        return
    
    # Nettoyer les données géographiques
    geo_data = data.dropna(subset=['latitude', 'longitude']).copy()
    
    if len(geo_data) == 0:
        st.warning("Aucune donnée géographique valide trouvée")
        return
    
    st.info(f"📍 {len(geo_data)} ménages géolocalisés sur {len(data)} au total")
    
    # Sélection du type de carte
    map_type = st.radio(
        "Type de visualisation",
        ["Carte des ménages", "Heatmap de densité", "Carte par province"],
        horizontal=True
    )
    
    if map_type == "Carte des ménages":
        # Carte avec marqueurs colorés selon le statut
        geo_data['statut'] = geo_data.apply(
            lambda row: 'Servi correctement' if row['indic_correct'] == 1 
            else 'Servi' if row['indic_servi'] == 1 
            else 'Non servi',
            axis=1
        )
        
        fig = px.scatter_mapbox(
            geo_data,
            lat='latitude',
            lon='longitude',
            color='statut',
            color_discrete_map={
                'Servi correctement': VisualizationEngine.COLOR_PALETTE['success'],
                'Servi': VisualizationEngine.COLOR_PALETTE['warning'],
                'Non servi': VisualizationEngine.COLOR_PALETTE['danger']
            },
            hover_data=['province', 'district', 'village'],
            zoom=6,
            height=600,
            title='<b>Carte des ménages enquêtés</b>'
        )
        
        fig.update_layout(mapbox_style="open-street-map")
        st.plotly_chart(fig, use_container_width=True)
    
    elif map_type == "Heatmap de densité":
        fig = px.density_mapbox(
            geo_data,
            lat='latitude',
            lon='longitude',
            z='indic_servi',
            radius=10,
            zoom=6,
            height=600,
            title='<b>Densité des ménages servis</b>'
        )
        
        fig.update_layout(mapbox_style="open-street-map")
        st.plotly_chart(fig, use_container_width=True)
    
    else:  # Carte par province
        province_centers = geo_data.groupby('province').agg({
            'latitude': 'mean',
            'longitude': 'mean',
            'indic_servi': 'mean',
            'province': 'count'
        }).reset_index()
        province_centers.columns = ['province', 'latitude', 'longitude', 'taux_couverture', 'total']
        province_centers['taux_couverture'] *= 100
        
        fig = px.scatter_mapbox(
            province_centers,
            lat='latitude',
            lon='longitude',
            size='total',
            color='taux_couverture',
            color_continuous_scale='RdYlGn',
            hover_name='province',
            hover_data={'total': True, 'taux_couverture': ':.1f'},
            zoom=5,
            height=600,
            title='<b>Taux de couverture par province</b>'
        )
        
        fig.update_layout(mapbox_style="open-street-map")
        st.plotly_chart(fig, use_container_width=True)


def page_statistics(data: pd.DataFrame):
    """Page avec statistiques avancées"""
    
    st.markdown("## 📊 Statistiques Avancées")
    
    # Statistiques descriptives
    st.markdown("### 📈 Statistiques descriptives")
    
    numeric_cols = data.select_dtypes(include=[np.number]).columns.tolist()
    selected_cols = st.multiselect(
        "Sélectionner les variables à analyser",
        numeric_cols,
        default=['nb_personnes', 'nb_milda_recues', 'nb_milda_attendues'][:len(numeric_cols)]
    )
    
    if selected_cols:
        desc_stats = data[selected_cols].describe().T
        desc_stats['missing'] = data[selected_cols].isnull().sum()
        desc_stats['missing_pct'] = (desc_stats['missing'] / len(data) * 100).round(2)
        
        st.dataframe(desc_stats, use_container_width=True)
    
    st.markdown("---")
    
    # Distributions
    st.markdown("### 📊 Distributions des variables")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if 'nb_personnes' in data.columns:
            fig = px.histogram(
                data,
                x='nb_personnes',
                nbins=30,
                title='<b>Distribution de la taille des ménages</b>',
                color_discrete_sequence=[VisualizationEngine.COLOR_PALETTE['primary']]
            )
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        if 'ecart_distribution' in data.columns:
            fig = px.histogram(
                data[data['menage_servi'] == 'Oui'],
                x='ecart_distribution',
                nbins=20,
                title='<b>Distribution des écarts de distribution</b>',
                color_discrete_sequence=[VisualizationEngine.COLOR_PALETTE['info']]
            )
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    st.markdown("---")
    
    # Corrélations (si scipy disponible)
    if STATS_AVAILABLE and len(selected_cols) > 1:
        st.markdown("### 🔗 Matrice de corrélation")
        
        corr_matrix = data[selected_cols].corr()
        
        fig = px.imshow(
            corr_matrix,
            text_auto='.2f',
            aspect="auto",
            color_continuous_scale='RdBu_r',
            title='<b>Corrélations entre variables</b>'
        )
        fig.update_layout(height=500)
        st.plotly_chart(fig, use_container_width=True)
    
    # Détection d'anomalies
    st.markdown("### 🔍 Détection d'anomalies")
    
    if 'nb_personnes' in data.columns:
    # 1. Créer une copie des données sans les valeurs manquantes pour cette colonne
    # Cela garantit que l'index de 'clean_data' sera le même que celui du masque 'outliers'
        clean_data = data.dropna(subset=['nb_personnes']).copy()
    
    # 2. Calculer les outliers sur ces données propres
        outliers = DataProcessor.detect_outliers(clean_data['nb_personnes'])
        n_outliers = outliers.sum()
    
        st.info(f"🔎 {n_outliers} valeurs aberrantes détectées dans la taille des ménages")
    
        if n_outliers > 0:
        # 3. Utiliser clean_data (et non data) pour le filtrage
            outlier_data = clean_data[outliers]
        
            st.dataframe(
            outlier_data[['province', 'district', 'village', 'nb_personnes', 'nb_milda_recues']].head(20),
            use_container_width=True
        )

from staticmap import StaticMap, CircleMarker
from PIL import Image

def get_village_map_screenshot(df_village, village_name, lat_col, lon_col):
    """
    Génère une image statique du village SANS Selenium.
    Ultra-léger pour Streamlit Cloud.
    """
    #st.text(f"📸 Génération statique : {village_name}...")
    
    # Nom du fichier image temporaire
    safe_name = "".join([c for c in village_name if c.isalnum()]).rstrip()
    temp_png = f"static_map_{safe_name}.png"

    try:
        # 1. Création de la carte statique
        # Comprendre le Zoom : staticmap utilise un zoom OSM (18 = très proche)
        # Pour le satellite gratuit, on utilise 'osm-satellite' si disponible ou OSM standard
        # Note: L'imagerie Airbus précise nécessite une API payante sans Selenium.
        # On utilise ici OSM standard pour la stabilité absolue.
        
        # Pour une vue satellite gratuite (Esri), décommentez la ligne url_template
        # Mais attention, OSM standard est plus rapide et stable sur le Cloud.
        # url_template = 'http://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}'
        
        m = StaticMap(1000, 800) # Taille de l'image finale en pixels

        # 2. Ajout de tous les points du village sur la carte
        for _, row in df_village.iterrows():
            if pd.notna(row[lat_col]) and pd.notna(row[lon_col]):
                
                is_servi = row.get('indic_servi') == 1
                color = '#00FF00' if is_servi else '#FF0000' # Vert/Rouge
                
                # Ajout du point (Marker) : Longitude, Latitude
                marker = CircleMarker((row[lon_col], row[lat_col]), color, 12)
                m.add_marker(marker)

        # 3. Calcul automatique du zoom idéal pour voir tous les points
        # staticmap le gère si on ne spécifie pas de zoom dans render()
        image = m.render() 
        
        # 4. Sauvegarde de l'image
        image.save(temp_png)
        
        if os.path.exists(temp_png):
            return temp_png
        else:
            return None

    except Exception as e:
        st.error(f"Erreur carte statique sur {village_name}: {str(e)}")
        # En cas d'erreur de tuiles, on crée une image blanche avec un message
        img = Image.new('RGB', (1000, 800), color = 'white')
        img.save(temp_png)
        return temp_png

def generate_visual_report(data, doc):
    """Génère le rapport Word avec images statiques (Méthode Survie)"""
    
    # Identification GPS
    lat_col = next((c for c in data.columns if 'lat' in c.lower()), None)
    lon_col = next((c for c in data.columns if 'lon' in c.lower() or 'lng' in c.lower()), None)
    
    if not lat_col or not lon_col:
        st.error("Coordonnées GPS manquantes.")
        return doc

    # Entête
    doc.add_heading('Rapport de Suivi Géospatial - CDM 2026', level=0)
    doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y')}").italic = True
    doc.add_paragraph("Cartographie : OpenStreetMap Statique | Échelle automatique").italic = True
    
    districts = [d for d in data['district'].unique() if pd.notna(d)]
    
    # Interface
    progress_bar = st.progress(0)
    status_text = st.empty()

    for idx, dist in enumerate(districts):
        status_text.info(f"📍 District : **{dist}** ({idx + 1}/{len(districts)})")
        progress_bar.progress((idx + 1) / len(districts))

        df_dist = data[data['district'] == dist]
        villages = [v for v in df_dist['village'].unique() if pd.notna(v)]
        
        # Sélection aléatoire (2 max)
        selected_villages = random.sample(list(villages), min(2, len(villages)))
        
        doc.add_heading(f'District Sanitaire : {dist}', level=1)

        for vil in selected_villages:
            df_vil = df_dist[df_dist['village'] == vil].dropna(subset=[lat_col, lon_col])
            
            if not df_vil.empty:
                doc.add_heading(f'Village : {vil}', level=2)
                
                # Génération statique (Très rapide, pas de crash RAM)
                img_path = get_village_map_screenshot(df_vil, vil, lat_col, lon_col)
                
                if img_path and os.path.exists(img_path):
                    # Insertion Word
                    doc.add_picture(img_path, width=Inches(5.8))
                    
                    # Légende
                    tot = len(df_vil)
                    serv = len(df_vil[df_vil['indic_servi'] == 1])
                    doc.add_paragraph(f"Carte de {vil}. Ménages : {tot} (Servis [Vert] : {serv})").italic = True
                    
                    # Nettoyage
                    os.remove(img_path)
            
            doc.add_page_break()

    progress_bar.empty()
    status_text.success("✅ Rapport généré !")
    return doc

def page_export(data: pd.DataFrame, tables: Dict[str, pd.DataFrame]):
    """Page d'export et de génération de rapports - Version Statique Stable"""
    
    st.markdown("## 📥 Export et Rapports")
    
    # --- 1. DÉTECTION GÉNÉRALE DES COLONNES GPS (CRITIQUE) ---
    # On définit ces variables au début pour qu'elles soient visibles par toutes les colonnes
    lat_col = next((c for c in data.columns if 'lat' in c.lower()), None)
    lon_col = next((c for c in data.columns if 'lon' in c.lower() or 'lng' in c.lower()), None)
    
    st.markdown("### 📊 Options d'export")
    
    # Calcul des métriques pour le rapport
    metrics = MetricsCalculator.calculate_coverage_metrics(data)
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    # --- COLONNE 1 : EXCEL ---
    with col1:
        st.markdown("#### Excel")
        if st.button("📊 Générer Excel", key="btn_excel", use_container_width=True):
            with st.spinner("Génération du rapport Excel..."):
                excel_file = ReportGenerator.generate_excel_report(data, tables, metrics)
                st.download_button(
                    label="⬇️ Télécharger Excel",
                    data=excel_file,
                    file_name=f"rapport_milda_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
    
    # --- COLONNE 2 : JSON ---
    with col2:
        st.markdown("#### JSON")
        if st.button("📋 Générer JSON", key="btn_json", use_container_width=True):
            json_report = ReportGenerator.generate_json_report(data, metrics)
            st.download_button(
                label="⬇️ Télécharger JSON",
                data=json_report,
                file_name=f"rapport_milda_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                use_container_width=True
            )
    
    # --- COLONNE 3 : CSV ---
    with col3:
        st.markdown("#### CSV")
        csv_data = data.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="⬇️ Télécharger CSV",
            data=csv_data,
            file_name=f"donnees_milda_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )

    # --- COLONNE 4 : GOOGLE EARTH (ZIP/KML) ---
    with col4:
        st.markdown("#### Google Earth")
        if st.button("🌍 Pack ZIP (KML)", key="btn_kml", use_container_width=True):
            if lat_col and lon_col:
                with st.spinner("Génération des fichiers KML..."):
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        df_geo = data.dropna(subset=[lat_col, lon_col])
                        grouped = df_geo.groupby(['province', 'district'])
                        
                        for (prov_name, dist_name), group in grouped:
                            kml = simplekml.Kml()
                            for _, row in group.iterrows():
                                pnt = kml.newpoint(name="") # Vide selon votre souhait
                                pnt.coords = [(row[lon_col], row[lat_col])]
                                is_servi = row.get('indic_servi') == 1
                                pnt.style.iconstyle.color = 'ff00ff00' if is_servi else 'ff0000ff'
                                pnt.style.iconstyle.icon.href = 'http://maps.google.com/mapfiles/kml/paddle/wht-blank.png'
                                pnt.description = (
                                f"<b>ID Ménage:</b> {row.get('id_menage', 'N/A')}<br>"
                                f"<b>Province:</b> {row.get('province', 'N/A')}<br>"
                                f"<b>District:</b> {row.get('district', 'N/A')}<br>"
                                f"<b>Centre de Santé:</b> {row.get('centre_sante', 'N/A')}<br>"
                                f"<b>Village:</b> {row.get('village', 'N/A')}<br>"
                                f"<b>Agent Enquêteur:</b> {row.get('agent_name', 'N/A')}"
                            )
                            
                            clean_prov = str(prov_name).replace("/", "-")
                            clean_dist = str(dist_name).replace("/", "-")
                            zip_file.writestr(f"{clean_prov}/{clean_dist}.kml", kml.kml())

                    st.download_button(
                        label="⬇️ Télécharger ZIP",
                        data=zip_buffer.getvalue(),
                        file_name=f"cartographie_{datetime.now().strftime('%Y%m%d')}.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
            else:
                st.error("Colonnes GPS introuvables.")

    # --- COLONNE 5 : RAPPORT VISUEL (WORD + STATICMAP) ---
    with col5:
        st.markdown("#### Visuels")
        if st.button("🖼️ Rapport Word Villages", key="btn_word", use_container_width=True):
            if not lat_col or not lon_col:
                st.error("Coordonnées GPS manquantes pour les visuels.")
            else:
                # --- INITIALISATION ---
                districts = [d for d in data['district'].unique() if pd.notna(d)]
                total_dist = len(districts)
                
                progress_bar = st.progress(0)
                status_text = st.empty() 
                
                with st.spinner("Génération des cartes statiques..."):
                    doc = Document()
                    doc.add_heading('Suivi Visuel par Village - CDM 2026', level=0)
                    doc.add_paragraph("Cartographie : OpenStreetMap Statique | Villages > 1 ménage").italic = True
                    
                    for idx, dist in enumerate(districts):
                        progression = (idx + 1) / total_dist
                        progress_bar.progress(progression)
                        status_text.info(f"📍 District {idx+1}/{total_dist} : **{dist}**")
                        
                        df_dist = data[data['district'] == dist]
                        
                        # --- FILTRAGE DES VILLAGES DE TAILLE > 1 ---
                        # On compte le nombre de ménages par village
                        counts = df_dist['village'].value_counts()
                        # On ne garde que les noms de villages qui apparaissent plus d'une fois
                        vils_eligibles = counts[counts > 10].index.tolist()
                        # On s'assure qu'ils ne sont pas nuls
                        vils_eligibles = [v for v in vils_eligibles if pd.notna(v)]
                        
                        if not vils_eligibles:
                            continue # Passer au district suivant si aucun village n'est assez grand
                        
                        # Sélection aléatoire parmi les villages éligibles
                        selected = random.sample(vils_eligibles, min(2, len(vils_eligibles)))
                        
                        doc.add_heading(f"District : {dist}", level=1)
                        
                        for vil_name in selected:
                            df_vil = df_dist[df_dist['village'] == vil_name].dropna(subset=[lat_col, lon_col])
                            
                            if not df_vil.empty:
                                doc.add_heading(f"Village : {vil_name}", level=2)
                                try:
                                    img_path = get_village_map_screenshot(df_vil, vil_name, lat_col, lon_col)
                                    
                                    if img_path and os.path.exists(img_path):
                                        doc.add_picture(img_path, width=Inches(5.5))
                                        
                                        servis = len(df_vil[df_vil['indic_servi'] == 1])
                                        doc.add_paragraph(f"Analyse de {len(df_vil)} ménages à {vil_name} (Servis : {servis}).")
                                        
                                        os.remove(img_path)
                                except Exception as e:
                                    doc.add_paragraph(f"Erreur image pour {vil_name} : {str(e)}")
                                
                                doc.add_page_break()
                    
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    
                    progress_bar.empty()
                    status_text.success(f"✅ Rapport généré !")
                    
                    st.download_button(
                        label="⬇️ Télécharger Word", 
                        data=doc_io.getvalue(), 
                        file_name=f"Rapport_Carto_MILDA_{datetime.now().strftime('%d_%m')}.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
    

    st.markdown("---")
    st.markdown("### 👁️ Prévisualisation des données")
    
    preview_option = st.selectbox(
        "Sélectionner un tableau à prévisualiser",
        ["Données brutes"] + list(tables.keys())
    )
    
    if preview_option == "Données brutes":
        st.dataframe(data.head(100), use_container_width=True)
        st.caption(f"Affichage des 100 premières lignes sur {len(data)} au total")
    else:
        st.dataframe(tables[preview_option], use_container_width=True)
    
    st.markdown("---")
    
    # Résumé des métriques
    st.markdown("### 📈 Résumé des métriques")
    
    summary_df = pd.DataFrame([metrics]).T
    summary_df.columns = ['Valeur']
    st.dataframe(summary_df, use_container_width=True)


import io

def page_agent_tracking(data: pd.DataFrame):
    st.markdown("## 🏃 Suivi du parcours des agents")

    df_track = data.copy()

    # --- Identification des colonnes ---
    start_col = 'start' if 'start' in df_track.columns else 'S0Q01_start'
    end_col   = 'heure_interview' if 'heure_interview' in df_track.columns else 'end'
    date_col  = 'date_enquete' if 'date_enquete' in df_track.columns else 'S0Q01'

    # ---------- HELPERS ----------
    def parse_single_datetime(val):
        if pd.isna(val) or str(val).strip().lower() in ['nan', 'none', 'nat', '']:
            return pd.NaT
        return pd.to_datetime(val, errors='coerce', dayfirst=True)

    def combine_date_time(row, target_col):
        raw_val = str(row[target_col]).strip()
        if raw_val.lower() in ['nan', 'none', 'nat', '']:
            return pd.NaT
        if '-' in raw_val or '/' in raw_val:
            return parse_single_datetime(raw_val)
        date_part = str(row[date_col]).split(' ')[0]
        try:
            return pd.to_datetime(f"{date_part} {raw_val}", errors='coerce', dayfirst=True)
        except:
            return pd.NaT

    def safe_dt_series(series):
        """Force une Series issue de apply() en datetime64[ns] propre."""
        return pd.to_datetime(series, errors='coerce', utc=False)

    # ── 1. FILTRES ───────────────────────────────────────────────────────────
    col_f1, col_f2, col_f3, col_f4 = st.columns(4)

    with col_f1:
        prov_list = ["Toutes"] + sorted(df_track['province'].dropna().unique().tolist())
        sel_prov = st.selectbox("📍 Province", prov_list)
        if sel_prov != "Toutes":
            df_track = df_track[df_track['province'] == sel_prov]

    with col_f2:
        dist_list = ["Tous"] + sorted(df_track['district'].dropna().unique().tolist())
        sel_dist = st.selectbox("🏙️ District", dist_list)
        if sel_dist != "Tous":
            df_track = df_track[df_track['district'] == sel_dist]

    with col_f3:
        vill_list = ["Tous"] + sorted(df_track['village'].dropna().unique().tolist())
        sel_vill = st.selectbox("🏡 Village", vill_list)
        if sel_vill != "Tous":
            df_track = df_track[df_track['village'] == sel_vill]

    with col_f4:
        t_dates = df_track[date_col].apply(parse_single_datetime).dropna()
        min_d = t_dates.min().date() if not t_dates.empty else datetime.now().date()
        max_d = t_dates.max().date() if not t_dates.empty else datetime.now().date()
        sel_date = st.date_input("📅 Date", value=(min_d, max_d))

        if isinstance(sel_date, tuple) and len(sel_date) == 2:
            df_track['temp_day'] = df_track[date_col].apply(parse_single_datetime).dt.date
            df_track = df_track[
                (df_track['temp_day'] >= sel_date[0]) &
                (df_track['temp_day'] <= sel_date[1])
            ]

    if df_track.empty:
        st.warning("Aucune donnée pour cette sélection.")
        return

    # ── 2. CONVERSION DATETIME (sécurisée) ───────────────────────────────────
    df_track['start_dt'] = safe_dt_series(
        df_track.apply(lambda r: combine_date_time(r, start_col), axis=1)
    )
    df_track['end_dt'] = safe_dt_series(
        df_track.apply(lambda r: combine_date_time(r, end_col), axis=1)
    )

    # ── 3. DURÉES ─────────────────────────────────────────────────────────────
    diff_td = pd.to_timedelta(df_track['end_dt'] - df_track['start_dt'], errors='coerce')
    df_track['Duree_min'] = diff_td.dt.total_seconds() / 60
    df_track.loc[df_track['Duree_min'] < 0, 'Duree_min'] = np.nan

    # ── 4. CARTE ──────────────────────────────────────────────────────────────
    df_map = df_track.dropna(subset=['latitude', 'longitude', 'agent_name', 'end_dt']).copy()

    if df_map.empty:
        st.warning("Aucune donnée GPS exploitable pour cette sélection.")
    else:
        # Garantie finale : end_dt est bien datetime avant .dt
        df_map['end_dt'] = safe_dt_series(df_map['end_dt'])
        df_map = df_map.sort_values(['agent_name', 'end_dt'])
        df_map['heure_texte'] = df_map['end_dt'].dt.strftime('%H:%M').fillna('--:--')

        agents = sorted(df_map['agent_name'].unique())
        c1, c2 = st.columns([2, 1])
        with c1:
            sel_agent = st.selectbox("👤 Enquêteur", agents)
        with c2:
            style = st.selectbox("🗺️ Fond", ["Satellite", "Rues", "Clair"])

        path = df_map[df_map['agent_name'] == sel_agent]

        if path.empty:
            st.info("Aucune donnée GPS pour cet enquêteur.")
        else:
            import plotly.express as px
            import plotly.graph_objects as go

            fig = px.line_mapbox(path, lat="latitude", lon="longitude", zoom=14, height=600)
            fig.add_trace(go.Scattermapbox(
                lat=path['latitude'],
                lon=path['longitude'],
                mode='markers+text',
                marker=dict(size=12, color='red'),
                text=path['heure_texte'],
                textposition="top right"
            ))

            mapbox_styles = {
                "Satellite": None,   # géré manuellement ci-dessous
                "Rues":  "open-street-map",
                "Clair": "carto-positron",
            }
            if style == "Satellite":
                fig.update_layout(
                    mapbox_style="white-bg",
                    mapbox_layers=[{
                        "sourcetype": "raster",
                        "source": [
                            "https://server.arcgisonline.com/ArcGIS/rest/services"
                            "/World_Imagery/MapServer/tile/{z}/{y}/{x}"
                        ]
                    }]
                )
            else:
                fig.update_layout(mapbox_style=mapbox_styles[style])

            fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})
            st.plotly_chart(fig, use_container_width=True)

    # ── 5. RAPPORT JOURNALIER ─────────────────────────────────────────────────
    st.divider()
    st.markdown("### 📋 Rapport journalier")

    def safe_strftime(val, fmt='%H:%M'):
        """Accepte un scalaire Timestamp ou NaT — jamais une Series."""
        try:
            if pd.isnull(val):
                return "N/A"
            return pd.Timestamp(val).strftime(fmt)
        except:
            return "N/A"
    
    rep = df_track.groupby('agent_name').agg(
        Enquêtes    = ('agent_name', 'count'),
        Moyenne_Min = ('Duree_min',  lambda x: round(x.mean(), 1) if not x.dropna().empty else 0),
        Début       = ('start_dt',   lambda x: safe_strftime(x.min())),
        Fin         = ('end_dt',     lambda x: safe_strftime(x.max())),
    ).reset_index()

    st.dataframe(
        rep.sort_values('Enquêtes', ascending=False),
        use_container_width=True,
        hide_index=True
    )

################################################################################
# 2. FONCTION page_data_quality() AMÉLIORÉE
################################################################################

def page_data_quality(data: pd.DataFrame):
    st.markdown("## 🛡️ Qualité des Données par Agent")
    
    if 'agent_name' not in data.columns:
        st.error("❌ Colonne 'agent_name' manquante.")
        return
    
    df_qc = data.copy()

    # --- 1. FONCTIONS DE CONVERSION ROBUSTES ---
    def get_safe_timestamp(row):
        """Combine date et heure de manière isolée pour éviter Mixed Inputs"""
        try:
            d_col = 'date_enquete' if 'date_enquete' in row else 'S0Q01'
            h_col = 'heure_interview' if 'heure_interview' in row else 'end'
            
            d_val = str(row[d_col]).split(' ')[0] if pd.notnull(row[d_col]) else ""
            h_val = str(row[h_col]).strip() if pd.notnull(row[h_col]) else ""
            
            if not d_val or h_val.lower() in ['nan', 'none', '']:
                return pd.NaT
            
            # Si l'heure contient déjà la date, on la parse directement
            if '-' in h_val or '/' in h_val:
                return pd.to_datetime(h_val, errors='coerce', dayfirst=True)
                
            return pd.to_datetime(f"{d_val} {h_val}", errors='coerce', dayfirst=True)
        except:
            return pd.NaT

    # --- 2. TRAITEMENT DES DONNÉES ---
    with st.spinner("Analyse de la qualité en cours..."):
        # Création du timestamp sécurisé
        df_qc['timestamp'] = df_qc.apply(get_safe_timestamp, axis=1)
        
        def calculate_agent_quality(agent_df):
            total = len(agent_df)
            if total == 0: return None
            
            # Complétude GPS (Latitude présente)
            comp_gps = (agent_df['latitude'].notna().sum()) / total * 100
            
            # Complétude des données (moyenne des cellules non vides)
            comp_data = agent_df.notna().mean(axis=1).mean() * 100
            
            # Doublons temporels et Vitesse
            doublons_pct = 0
            vitesse_h = 0
            valid_ts = agent_df.dropna(subset=['timestamp'])
            
            if len(valid_ts) > 1:
                # Calcul des doublons sur le timestamp
                doublons_count = valid_ts['timestamp'].duplicated().sum()
                doublons_pct = (doublons_count / total) * 100
                
                # Calcul de la vitesse (enquêtes par heure)
                diff_time = (valid_ts['timestamp'].max() - valid_ts['timestamp'].min()).total_seconds() / 3600
                if diff_time > 0.1: # Éviter division par presque zéro
                    vitesse_h = total / diff_time

            # Score final pondéré
            # 50% Complétude, 30% GPS, 20% Absence de doublons
            score = (comp_data * 0.5) + (comp_gps * 0.3) + ((100 - doublons_pct) * 0.2)
            
            return {
                'Agent': agent_df['agent_name'].iloc[0],
                'Enquêtes': total,
                'Complétude (%)': round(comp_data, 1),
                'GPS (%)': round(comp_gps, 1),
                'Doublons (%)': round(doublons_pct, 1),
                'Vitesse (Enq/h)': round(vitesse_h, 2),
                'Score Qualité': round(score, 1)
            }

        # Calcul pour chaque agent
        quality_results = []
        unique_agents = df_qc['agent_name'].dropna().unique()
        for agent in unique_agents:
            res = calculate_agent_quality(df_qc[df_qc['agent_name'] == agent])
            if res: quality_results.append(res)
        
        if not quality_results:
            st.warning("Aucune donnée disponible pour calculer la qualité.")
            return

        quality_df = pd.DataFrame(quality_results).sort_values('Score Qualité', ascending=False)

    # --- 3. AFFICHAGE DES RÉSULTATS ---
    import plotly.express as px
    
    # Cartes de scores globaux
    avg_score = quality_df['Score Qualité'].mean()
    m1, m2, m3 = st.columns(3)
    m1.metric("Score Moyen Global", f"{avg_score:.1f} / 100")
    m2.metric("Meilleure Performance", f"{quality_df['Score Qualité'].max():.1f}")
    m3.metric("Agents Analysés", len(quality_df))

    # Graphique de classement
    fig = px.bar(
        quality_df, 
        x='Agent', 
        y='Score Qualité', 
        color='Score Qualité',
        color_continuous_scale='RdYlGn',
        range_color=[0, 100],
        title="Classement de la Qualité par Agent"
    )
    st.plotly_chart(fig, use_container_width=True)

    # Tableau détaillé avec coloration
    st.markdown("### 📋 Détails des indicateurs de contrôle")
    
    # Style pour mettre en évidence les bons et mauvais scores
    styled_df = quality_df.style.background_gradient(subset=['Score Qualité'], cmap='RdYlGn', vmin=0, vmax=100) \
                                .background_gradient(subset=['Doublons (%)'], cmap='Reds', vmin=0, vmax=10)
    
    st.dataframe(styled_df, use_container_width=True, hide_index=True)
################################################################################
# 3. GÉNÉRATION AUTOMATIQUE DE RAPPORT (Format Word)
################################################################################

def create_table(document, data, headers):
    """Crée un tableau formaté dans le document Word"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # En-têtes
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Formater l'en-tête en gras
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Données
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    return table


def add_matplotlib_chart(document, data_series, title, chart_type='bar'):
    """Génère un graphique avec étiquettes de pourcentage et l'insère dans Word"""
    # Utiliser un style propre
    plt.style.use('seaborn-v0_8-muted') 
    fig, ax = plt.subplots(figsize=(7, 4))
    
    if chart_type == 'bar':
        # Création des barres
        data_series.plot(kind='bar', ax=ax, color='#2c3e50', edgecolor='white')
        ax.set_ylabel('')
        ax.set_xlabel('')
        
        # AJOUT DES LIBELLÉS (Valeurs au-dessus des barres)
        for p in ax.patches:
            height = p.get_height()
            ax.annotate(f'{height:.1f}%', 
                        (p.get_x() + p.get_width() / 2., height), 
                        ha='center', va='center', 
                        xytext=(0, 8), 
                        textcoords='offset points',
                        fontsize=9, fontweight='bold', color='#c0392b')
            
        # Ajuster l'échelle pour ne pas couper les étiquettes
        ax.set_ylim(0, max(data_series.max() * 1.2, 100))

    elif chart_type == 'pie':
        # Graphique en secteurs avec pourcentages intégrés
        data_series.plot(kind='pie', ax=ax, autopct='%1.1f%%', 
                         startangle=140, colors=['#27ae60', '#e67e22', '#3498db'])
        ax.set_ylabel('')

    plt.title(title, pad=20, fontsize=11, fontweight='bold')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    # Conversion en image pour Word
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    
    # Ajout au document Word
    document.add_picture(img_stream, width=Inches(5.5))
    last_p = document.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def calculate_milda_requis_custom(nb_personnes):
    """Calcule le nombre requis selon votre nouvelle table"""
    if nb_personnes <= 2: return 1
    elif nb_personnes <= 4: return 2
    elif nb_personnes >= 5: return 3  # 5-6 pers = 3, et 7+ pers = 3
    return 0

def add_custom_color_chart(document, data_series, title):
    plt.figure(figsize=(7, 4))
    
    # Définition des couleurs selon la valeur de l'index (la différence)
    colors = []
    for val in data_series.index:
        if val < 0: colors.append('#e74c3c') # Rouge pour sous-distribution
        elif val == 0: colors.append('#27ae60') # Vert pour conforme
        else: colors.append('#3498db') # Bleu pour sur-distribution
        
    ax = data_series.plot(kind='bar', color=colors, edgecolor='black')
    
    # Ajout des pourcentages au-dessus des barres
    for p in ax.patches:
        ax.annotate(f"{p.get_height():.1f}%", 
                    (p.get_x() + p.get_width() / 2., p.get_height()),
                    ha='center', va='center', xytext=(0, 8), 
                    textcoords='offset points', fontsize=9, fontweight='bold')

    plt.title(title, pad=20, fontweight='bold')
    plt.ylabel('Fréquence (%)')
    plt.ylim(0, max(data_series.max() * 1.2, 100))
    plt.tight_layout()

    # Sauvegarde et insertion
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close()
    img_stream.seek(0)
    document.add_picture(img_stream, width=Inches(5.5))

def add_custom_diff_chart(document, data, title):
    """Génère et insère le graphique coloré des différences dans le Word"""
    import matplotlib.pyplot as plt
    
    # 1. Préparation des données
    # On trie par index pour avoir -2, -1, 0, 1...
    stats = (data['diff_custom'].value_counts(normalize=True).sort_index() * 100)
    
    if stats.empty:
        return

    # 2. Création de la figure
    fig, ax = plt.subplots(figsize=(7, 4))
    
    # Définition des couleurs
    colors = []
    for val in stats.index:
        if val < 0: colors.append('#e74c3c')    # Rouge (Manque)
        elif val == 0: colors.append('#27ae60') # Vert (Conforme)
        else: colors.append('#3498db')          # Bleu (Surplus)
        
    bars = ax.bar(stats.index.astype(str), stats.values, color=colors, edgecolor='black')
    
    # 3. Ajout des libellés de pourcentage
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{height:.1f}%',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 5), textcoords="offset points",
                    ha='center', va='bottom', fontsize=9, fontweight='bold')

    # Mise en forme
    plt.title(title, pad=20, fontweight='bold')
    plt.ylabel('')
    plt.xlabel('Écart (Nombre de MILDA)')
    ax.set_ylim(0, max(stats.values) * 1.2) # Espace pour les étiquettes
    plt.tight_layout()

    # 4. Sauvegarde temporaire et insertion
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    
    document.add_picture(img_stream, width=Inches(5.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def add_chart_placeholder(document, title):
    """Ajoute un espace réservé pour un graphique"""
    p = document.add_paragraph()
    p.add_run(f"[GRAPHIQUE: {title}]").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_automatic_report(data: pd.DataFrame, tables: dict) -> io.BytesIO:
    """
    Génère un rapport automatique au format Word
    Structure inspirée de Analyse_denombrement_pilote.docx
    
    Returns:
        io.BytesIO: Document Word en mémoire
    """
    
    if not DOCX_AVAILABLE:
        st.error("❌ Bibliothèque python-docx non disponible. Installez-la avec: pip install python-docx")
        return None
    data = data.copy()
    #data_1 = data.copy()
    #data = data[data['consentement'] == 'Oui']
    
    # Créer le document
    doc = Document()
    
    # ========== PAGE DE TITRE ==========
    doc.add_heading('Analyse du dénombrement-distribution MILDA', 0)
    doc.add_heading('Campagne de Distribution de Masse 2026', level=2)
    
    p = doc.add_paragraph()
    p.add_run(f'Rapport généré le : {datetime.now().strftime("%d/%m/%Y à %H:%M")}\n').bold = True
    p.add_run(f'Période d\'analyse : ')
    if 'date_enquete' in data.columns:
        date_min = data['date_enquete'].min().strftime('%d/%m/%Y')
        date_max = data['date_enquete'].max().strftime('%d/%m/%Y')
        p.add_run(f'{date_min} au {date_max}')
    
    doc.add_page_break()
    
    # ========== CARACTÉRISTIQUES DES MÉNAGES ==========
    doc.add_heading('Profil du Chef de Ménage', level=1)
    
    # Tableau 1: Proportion des chefs de ménage
    doc.add_heading('Tableau : Proportion des chefs des ménages enquêtés', level=2)
    
    if 'menage_chef' in data.columns or any('chef' in col.lower() for col in data.columns):
        # Trouver la colonne appropriée
        chef_col = next((col for col in data.columns if 'chef' in col.lower()), None)
        if chef_col:
            chef_data = data[chef_col].value_counts()
            total = len(data)
            
            table_data = []
            for value, count in chef_data.items():
                freq = round(count / total * 100, 2)
                table_data.append([value, count, freq])
            table_data.append(['Total', count, 100.00])
            
            create_table(doc, table_data, ['Êtes-vous le Chef de ce ménage ?', 'Effectif', 'Fréquence'])
            doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True

    # Tableau Sexe
    if 'sexe' in data.columns:
        doc.add_heading('Tableau : Répartition des chefs de ménage par sexe', level=2)
        sexe_counts = data['sexe'].value_counts()
        total_s = len(data)
        table_sexe = [[v, c, f"{(c/total_s*100):.1f}"] for v, c in sexe_counts.items()]
        table_sexe.append(['Total', total_s, '100'])
        create_table(doc, table_sexe, ['Sexe', 'Effectif', 'Fréquence (%)'])
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True

    # Tableau Activité (activ_rev)
    if 'activ_rev' in data.columns:
        doc.add_heading('Tableau : Activité principale du chef de ménage', level=2)
        act_counts = data['activ_rev'].value_counts()
        total_a = len(data)
        table_act = [[v, c, f"{(c/total_a*100):.1f}"] for v, c in act_counts.items()]
        table_act.append(['Total', total_a, '100'])
        create_table(doc, table_act, ['Activité', 'Effectif', 'Fréquence (%)'])
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True


    
    # ========== SECTION : ANALYSE DE LA DISTRIBUTION ==========
    doc.add_heading('Analyse de la conformité de la distribution', level=1)
    
    # Application de votre règle personnalisée (7+ = 3 MILDA)
    def calculate_milda_requis_custom(nb_personnes):
        if nb_personnes <= 2: return 1
        elif nb_personnes <= 4: return 2
        elif nb_personnes >= 5: return 3
        return 0

    data['requis_custom'] = data['nb_personnes'].apply(calculate_milda_requis_custom)
    data['diff_custom'] = data['requis_custom'] - data['nb_milda_recues']

    # --- AJOUT DU GRAPHIQUE ---
    doc.add_paragraph("Le graphique ci-dessous présente la répartition des écarts constatés :")
    add_custom_diff_chart(doc, data, "Répartition des écarts de distribution par rapport à la norme")

    # --- AJOUT DU TABLEAU (votre structure demandée) ---
    diff_counts = data['diff_custom'].value_counts().sort_index()
    total_diff = len(data)
    
    table_rows = []
    for val, count in diff_counts.items():
        freq = (count / total_diff * 100)
        table_rows.append([int(val), count, f"{freq:.1f}"])
    
    table_rows.append(['Total', total_diff, '100'])
    create_table(doc, table_rows, ['Nombre de Différence', 'Effectif', 'Fréquence (%)'])
    doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True

    doc.add_page_break()
            
    def add_executive_summary(doc, data):
        doc.add_heading('RÉSUMÉ EXÉCUTIF : ALERTES DE PERFORMANCE', level=1)
        
        # 1. Calcul des indicateurs par Centre de Santé (National)
        cs_performance = data.groupby(['province', 'district']).agg(
            total=('indic_servi', 'count'),
            servis=('indic_servi', 'sum'),
            marques=('indic_marque', 'sum')
        ).reset_index()
    
        cs_performance['taux_couverture'] = (cs_performance['servis'] / cs_performance['total']) * 100
        cs_performance['taux_marquage'] = (cs_performance['marques'] / cs_performance['servis']) * 100
    
        # 2. Identification des zones critiques (Seuil < 80%)
        alerte_couverture = cs_performance[cs_performance['taux_couverture'] < 80].sort_values('taux_couverture')
        alerte_marquage = cs_performance[cs_performance['taux_marquage'] < 80].sort_values('taux_marquage')
    
        # --- SECTION : ALERTES COUVERTURE ---
        doc.add_heading('🚨 Zones à faible taux de couverture (< 80%)', level=2)
        if not alerte_couverture.empty:
            p = doc.add_paragraph(f"Les {len(alerte_couverture)} centres de santé suivants présentent une couverture insuffisante. Une vérification logistique ou une supervision de proximité est recommandée.")
            
            table_alert = [['Province', 'District Sanitaire', 'Taux Couverture']]
            for _, row in alerte_couverture.head(10).iterrows(): # Top 10 des plus critiques
                table_alert.append([
                    row['province'],
                    row['district'],
                    f"{row['taux_couverture']:.1f}%"
                ])
            create_table(doc, table_alert, table_alert[0]) # Utilise votre fonction create_table
        else:
            doc.add_paragraph("Félicitations : Tous les district dépassent 80% de couverture.")
    
        # --- SECTION : ALERTES MARQUAGE ---
        doc.add_heading('Zones à faible taux de marquage (< 80%)', level=2)
        if not alerte_marquage.empty:
            doc.add_paragraph("Le marquage des ménages est essentiel pour le suivi. Les zones suivantes sont en dessous des standards de qualité :")
            
            table_m = [['Province', 'District', 'Taux Marquage']]
            for _, row in alerte_marquage.head(10).iterrows():
                table_m.append([
                    row['province'],
                    row['district'],
                    f"{row['taux_marquage']:.1f}%"
                ])
            create_table(doc, table_m, table_m[0])
        else:
            doc.add_paragraph("Qualité technique : Le marquage est conforme aux attentes dans toutes les zones.")

    def add_provincial_dashboard(doc, data):
        doc.add_heading('PRINCIPAUX INDICATEURS PAR PROVINCE', level=1)
        doc.add_paragraph("Ce tableau compare la performance globale de chaque province pour l'ensemble des indicateurs clés de la CDM-2026.")
    
        # 1. Agrégation correcte des données par Province
        prov_stats = data.groupby('province').agg(
            nb_menages=('indic_servi', 'count'),   # Modifié : 'count' compte les lignes
            servis=('indic_servi', 'sum'),         # Somme des 1 pour les ménages servis
            marques=('indic_marque', 'sum'),       # Somme des 1 pour les ménages marqués
            corrects=('indic_correct', 'sum')      # Modifié : 'sum' pour additionner les ménages conformes
        ).reset_index()
    
        # 2. Calcul des indicateurs de performance (sécurisé)
        prov_stats['% Couverture'] = (prov_stats['servis'] / prov_stats['nb_menages'] * 100).round(1)
        prov_stats['% Marquage'] = (prov_stats['marques'] / prov_stats['servis'].replace(0, np.nan)).astype(float).round(1).fillna(0)
        prov_stats['% Qualité (Correct)'] = (prov_stats['corrects'] / prov_stats['servis'].replace(0, np.nan)).astype(float).round(1).fillna(0)
    
        # Tri par performance de couverture (du meilleur au moins bon)
        prov_stats = prov_stats.sort_values('% Couverture', ascending=False)
    
        # 3. Préparation des données pour le tableau Word
        table_headers = [
            'Province', 
            'Ménages Dénombrés', 
            '% Couverture', 
            '% Marquage', 
            '% Qualité'
        ]
        
        table_data = []
        for _, row in prov_stats.iterrows():
            table_data.append([
                str(row['province']),
                f"{int(row['nb_menages']):,}".replace(',', ' '), # Formatage espace pour milliers
                f"{row['% Couverture']}%",
                f"{row['% Marquage']}%",
                f"{row['% Qualité (Correct)']}%"
            ])
    
        # 4. Ajout d'une ligne de TOTAL NATIONAL
        total_n = prov_stats['nb_menages'].sum()
        total_s = prov_stats['servis'].sum()
        total_m = prov_stats['marques'].sum()
        total_c = prov_stats['corrects'].sum()
    
        table_data.append([
            'TOTAL NATIONAL',
            f"{int(total_n):,}".replace(',', ' '),
            f"{round(100 * total_s / total_n, 1)}%" if total_n > 0 else "0%",
            f"{round(100 * total_m / total_s, 1)}%" if total_s > 0 else "0%",
            f"{round(100 * total_c / total_n, 1)}%" if total_s > 0 else "0%"
        ])
    
        # 5. Création du tableau dans le document Word
        create_table(doc, table_data, table_headers)
          
        doc.add_paragraph("Note : Le % Qualité représente la proportion de ménages servis ayant reçu la MILDA conformément aux procédures standards.").italic = True
        # --- AJOUT DU GRAPHIQUE DE COMPARAISON ---
        doc.add_heading('Comparaison visuelle de la couverture par Province (%)', level=2)
        
        # On prépare les données pour le graphique
        # On utilise les noms de provinces et les taux de couverture calculés précédemment
        series_couverture = prov_stats.set_index('province')['% Couverture']
        
        # Appel de votre fonction de graphique (en mode 'bar' vertical)
        # Note : Assurez-vous que votre fonction add_matplotlib_chart gère bien les index
        add_matplotlib_chart(
            doc, 
            series_couverture, 
            'Taux de Couverture MILDA (%)', 
            'bar'
        )
        
        doc.add_paragraph("Figure 1 : Classement des provinces par taux de couverture décroissant.").italic = True

    #add_executive_summary(doc, data)
    add_provincial_dashboard(doc, data)

    # 1. Identifier toutes les provinces uniques
    if 'province' in data.columns:
        provinces = sorted(data['province'].dropna().unique()) # Notez le 's' à provinces
    else:
        st.error("La colonne 'province' est manquante.")
        provinces = []

    for prov in provinces:
        # Filtrer les données pour la province actuelle
        df_prov = data[data['province'] == prov].copy()
        
        # Titre de la section Province
        doc.add_heading(f'Analyse : Province de {prov}', level=1)
        
        # ==========================================================
        # NOUVELLE SECTION : ANALYSE PAR DISTRICT SANITAIRE
        # ==========================================================
        if 'district' in df_prov.columns and not df_prov.empty:
            doc.add_heading(f'Indicateurs par District Sanitaire - {prov}', level=2)
            
            # 1. Agrégation par District
            dist_stats = df_prov.groupby('district').agg(
                total=('district', 'count'),
                servis=('indic_servi', 'sum'),
                correct=('indic_correct', 'sum')
            ).reset_index()
            
            # 2. Calcul des indicateurs
            dist_stats['pct_servis'] = (100 * dist_stats['servis'] / dist_stats['total']).astype(float).round(1)
            dist_stats['pct_correct'] = (100 * dist_stats['correct'] / dist_stats['servis'].replace(0, np.nan)).astype(float).round(1).fillna(0)
            
            # 3. Graphique de couverture par District
            add_matplotlib_chart(doc, dist_stats.set_index('district')['pct_servis'], f'Couverture MILDA par District - {prov} (%)', 'bar')
            
            # 4. Tableau par District
            table_dist_data = []
            for _, row in dist_stats.iterrows():
                table_dist_data.append([
                    str(row['district']),
                    int(row['total']),
                    int(row['servis']),
                    f"{row['pct_servis']}%",
                    int(row['correct']),
                    f"{row['pct_correct']}%"
                ])
            
            create_table(doc, table_dist_data, [
                'District Sanitaire',
                'Ménages dénombrés',
                'Ménages servis',
                '% servis',
                'Correctement servis',
                '% correct'
            ])
            doc.add_paragraph(f'Source : Analyse par district pour la province de {prov}').italic = True
            doc.add_page_break()
        
        
        doc.add_heading(f'Analyse du marquage des ménages - {prov}', level=2)

        # --- ANALYSE DU MARQUAGE PAR DISTRICT ---
        doc.add_heading(f'Analyse du marquage par District Sanitaire - {prov}', level=2)
        
        if 'district' in df_prov.columns and not df_prov.empty:
            # Filtrer uniquement les ménages servis pour le calcul du marquage
            df_servis_dist = df_prov[df_prov['indic_servi'] == 1]
            
            if not df_servis_dist.empty:
                # 1. Calcul des statistiques par District
                dist_marquage_stats = df_servis_dist.groupby('district').agg(
                    servis=('indic_servi', 'count'),
                    marques=('indic_marque', 'sum')
                ).reset_index()
                
                # Calcul du pourcentage sécurisé (float)
                dist_marquage_stats['pct_marques'] = (100 * dist_marquage_stats['marques'] / dist_marquage_stats['servis']).astype(float).round(1)
                
                # 2. Graphique par District
                add_matplotlib_chart(
                    doc, 
                    dist_marquage_stats.set_index('district')['pct_marques'], 
                    f'Taux de marquage par District - {prov} (%)', 
                    'bar'
                )
                
                # 3. Tableau par District
                table_dist_marquage = []
                for _, row in dist_marquage_stats.iterrows():
                    table_dist_marquage.append([
                        str(row['district']),
                        int(row['servis']),
                        int(row['marques']),
                        f"{row['pct_marques']}%"
                    ])
                
                # Ajouter la ligne de Total Province pour le marquage
                t_servis = dist_marquage_stats['servis'].sum()
                t_marques = dist_marquage_stats['marques'].sum()
                t_pct = round(100 * t_marques / t_servis, 1) if t_servis > 0 else 0
                
                table_dist_marquage.append(['TOTAL PROVINCE', t_servis, t_marques, f"{t_pct}%"])
                
                create_table(doc, table_dist_marquage, [
                    'District Sanitaire', 
                    'Ménages servis', 
                    'Ménages marqués', 
                    '% marqués'
                ])
            else:
                doc.add_paragraph("Aucun ménage servi enregistré pour l'analyse du marquage dans ces districts.")
        
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True
        
    
    # ========== ANALYSE DE LA DISTRIBUTION ==========
    doc.add_heading('Analyse de la distribution des moustiquaires', level=1)
    
    # Calcul des écarts
    if 'ecart_distribution' in data.columns:
        distribution_data = data[data['menage_servi'] == 'Oui'].copy()
        
        moins_norme = (distribution_data['ecart_distribution'] < 0).sum()
        norme_ok = (distribution_data['ecart_distribution'] == 0).sum()
        plus_norme = (distribution_data['ecart_distribution'] > 0).sum()
        total_dist = len(distribution_data)
        
        pct_moins = round(100 * moins_norme / total_dist, 1) if total_dist > 0 else 0
        pct_ok = round(100 * norme_ok / total_dist, 1) if total_dist > 0 else 0
        pct_plus = round(100 * plus_norme / total_dist, 1) if total_dist > 0 else 0
        
        p = doc.add_paragraph()
        p.add_run(
            f'Il ressort que {pct_moins}% des ménages ont reçu des moustiquaires en moins selon la norme prévue '
            f'et {pct_plus}% ont reçu des moustiquaires en plus que ce qui était prévu. '
        )
        
        doc.add_heading('Tableau : Répartition selon la norme', level=2)
        
        table_data = [
            ['Moins que la norme', moins_norme, pct_moins],
            ['Norme respectée', norme_ok, pct_ok],
            ['Plus que la norme', plus_norme, pct_plus],
            ['Total', total_dist, 100.0]
        ]
        
        create_table(doc, table_data, ['Nombre des moustiquaires reçues', 'Effectif', 'Fréquence (%)'])
        
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True
    
    #doc.add_page_break()
    
    # ========== SECTION : SENSIBILISATION (INFORMATION) ==========
    doc.add_heading('Information sur la campagne de distribution des MILDA', level=1)
    
    # --- 1. TABLEAU DE DISTRIBUTION PAR PROVINCE (SYNTHÈSE NATIONALE) ---
    if 'province' in data.columns:
        doc.add_heading('Synthèse d\'Information sur la campagne de distribution des MILDA par Province', level=2)
        
        prov_sensi = data.groupby('province').agg(
            total=('province', 'count'),
            informes=('indic_info', 'sum')
        ).reset_index()
        
        # Calcul du pourcentage par province
        prov_sensi['pct_informes'] = (100 * prov_sensi['informes'] / prov_sensi['total']).astype(float).round(1)
        
        table_prov_data = []
        for _, row in prov_sensi.iterrows():
            table_prov_data.append([
                str(row['province']),
                int(row['total']),
                int(row['informes']),
                f"{row['pct_informes']}%"
            ])
        
        # Ligne de Total National
        t_n_total = prov_sensi['total'].sum()
        t_n_informes = prov_sensi['informes'].sum()
        t_n_pct = round(100 * t_n_informes / t_n_total, 1) if t_n_total > 0 else 0
        table_prov_data.append(['TOTAL NATIONAL', t_n_total, t_n_informes, f"{t_n_pct}%"])
        
        create_table(doc, table_prov_data, [
            'Province', 
            'Ménages total', 
            'Ménages informés', 
            '% informés'
        ])
        doc.add_paragraph('Source : Données consolidées CDM-2026').italic = True
        #doc.add_page_break()
    
    # ========== SECTION : SENSIBILISATION (INFORMATION) ==========
    #doc.add_heading('Information sur l\'utilisation correcte des MILDA', level=1)
    
    if 'province' in data.columns:
        provinces = sorted(data['province'].dropna().unique())
        
        for prov in provinces:
            df_prov = data[data['province'] == prov].copy()
            
            # --- 1. Titre de la Province ---
            doc.add_heading(f'Province : {prov}', level=2)
            
            # --- 2. Analyse par District Sanitaire au sein de la Province ---
            if 'district' in df_prov.columns and not df_prov.empty:
                doc.add_heading(f'Information sur la campagne de distribution des MILDA par District Sanitaire - {prov}', level=3)
                
                dist_sensi = df_prov.groupby('district').agg(
                    total=('district', 'count'),
                    informes=('indic_info', 'sum')
                ).reset_index()
                
                dist_sensi['pct_informes'] = (100 * dist_sensi['informes'] / dist_sensi['total']).astype(float).round(1)
                
                table_dist_data = []
                for _, row in dist_sensi.iterrows():
                    table_dist_data.append([
                        str(row['district']),
                        int(row['total']),
                        int(row['informes']),
                        f"{row['pct_informes']}%"
                    ])
                
                # Total pour la province (ligne de résumé)
                t_total_p = dist_sensi['total'].sum()
                t_informes_p = dist_sensi['informes'].sum()
                t_pct_p = round(100 * t_informes_p / t_total_p, 1) if t_total_p > 0 else 0
                table_dist_data.append(['TOTAL PROVINCE', t_total_p, t_informes_p, f"{t_pct_p}%"])
                
                create_table(doc, table_dist_data, ['District Sanitaire', 'Ménages total', 'Ménages informés', '% informés'])
                doc.add_paragraph(f'Source : CDM-2026 - Province de {prov}').italic = True
    
    # ========== SECTION : SENSIBILISATION LORS DE LA CAMPAGNE ==========
    doc.add_heading('Sensibilisation sur l\'utilisation correcte lors de la campagne', level=1)
    
    info_col = 'sensibilise'
    
    if info_col in data.columns:
        # Nettoyage : On s'assure que la colonne est traitée comme du texte pour la comparaison
        df_clean = data.copy()
        df_clean[info_col] = df_clean[info_col].astype(str).str.strip().str.lower()
    
        # --- 1. SYNTHÈSE NATIONALE PAR PROVINCE ---
        doc.add_heading('Tableau : Proportion des ménages sensibilisés par Province', level=2)
        
        # Définition des valeurs considérées comme "Vrai/Oui"
        # On accepte 'oui', '1', '1.0', 'yes'
        valeurs_positives = ['oui', '1', '1.0', 'yes']
        
        prov_sensi_global = df_clean.groupby('province').agg(
            total=(info_col, 'count'),
            sensibilises=(info_col, lambda x: x.isin(valeurs_positives).sum())
        ).reset_index()
        
        prov_sensi_global['pct'] = (100 * prov_sensi_global['sensibilises'] / prov_sensi_global['total']).astype(float).round(1)
        
        table_prov = []
        for _, row in prov_sensi_global.iterrows():
            table_prov.append([str(row['province']), int(row['total']), int(row['sensibilises']), f"{row['pct']}%"])
        
        # Ligne Total National
        t_n_total = prov_sensi_global['total'].sum()
        t_n_sensi = prov_sensi_global['sensibilises'].sum()
        t_n_pct = round(100 * t_n_sensi / t_n_total, 1) if t_n_total > 0 else 0
        table_prov.append(['TOTAL NATIONAL', t_n_total, t_n_sensi, f"{t_n_pct}%"])
        
        create_table(doc, table_prov, ['Province', 'Ménages total', 'Ménages sensibilisés', '% sensibilisés'])
        doc.add_page_break()
    
        # --- 2. DÉTAIL PAR PROVINCE > DISTRICT > CS ---
        provinces = sorted(df_clean['province'].dropna().unique())
        
        for prov in provinces:
            df_prov = df_clean[df_clean['province'] == prov].copy()
            doc.add_heading(f'Analyse détaillée : Province de {prov}', level=2)
            
            # --- A. PAR DISTRICT ---
            if 'district' in df_prov.columns:
                doc.add_heading(f'Sensibilisation par District Sanitaire - {prov}', level=3)
                dist_stats = df_prov.groupby('district').agg(
                    total=(info_col, 'count'),
                    sensi=(info_col, lambda x: x.isin(valeurs_positives).sum())
                ).reset_index()
                
                dist_stats['pct'] = (100 * dist_stats['sensi'] / dist_stats['total']).astype(float).round(1)
                table_dist = [[str(r['district']), int(r['total']), int(r['sensi']), f"{r['pct']}%"] for _, r in dist_stats.iterrows()]
                create_table(doc, table_dist, ['District', 'Ménages total', 'Ménages sensibilisés', '%'])

    # Source d'information (source)
    if 'source' in data.columns:
        doc.add_heading('Tableau : Sources d\'information sur la campagne', level=2)
        # Gestion du choix multiple
        source_series = data['source'].str.split(', ').explode().value_counts()
        total_resp = len(data)
        table_source = [[v, c, f"{(c/total_resp*100):.1f}"] for v, c in source_series.items()]
        create_table(doc, table_source, ['Source citée', 'Effectif', '% de ménages'])
        doc.add_paragraph("Note: Un ménage peut citer plusieurs sources.").italic = True
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True

    # Instructions reçues (conseil)
    if 'conseil' in data.columns:
        doc.add_heading('Tableau : Instructions d\'utilisation reçues (Conseils)', level=2)
        conseil_series = data['conseil'].str.split(', ').explode().value_counts()
        table_conseil = [[v, c, f"{(c/total_resp*100):.1f}"] for v, c in conseil_series.items()]
        create_table(doc, table_conseil, ['Conseil prodigué', 'Effectif', '% de ménages'])
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True

    # ========== CALCUL DES COMPTEURS DE SCANS (À ajouter au début de la fonction) ==========

    # 1. Vérifiez d'abord si la colonne brute existe (souvent nommée 'scan_milda' ou '_scan_milda_count')
    # Si KoBo n'envoie pas le compte, nous allons le déduire des données
    if 'id_scan' in data.columns:
        # Si id_scan est une liste ou contient les données du repeat
        data['id_scan_count'] = data['id_scan'].apply(lambda x: len(x) if isinstance(x, list) else (1 if pd.notnull(x) else 0))
    else:
        # Si la colonne n'existe pas du tout, on initialise à 0 pour éviter le crash
        # Note : Vérifiez dans votre export Excel le nom exact de la colonne du repeat
        data['id_scan_count'] = 0 
    
    # ========== ANALYSE DES SCANS QR (La section qui posait erreur) ==========
    doc.add_heading('Analyse du scannage des codes QR', level=1)
    
    # On utilise nb_milda_recues (Nombre de MILDA reçu selon le répondant )
    total_recues = data['nb_milda_recues'].astype(float).sum() if 'nb_milda_recues' in data.columns else 0
    nb_scannes = data['id_scan_count'].sum() 
    nb_non_scannes = max(0, total_recues - nb_scannes)
    
    # Affichage du Tableau
    table_scan = [
        ["Nombre des scans QR", int(nb_scannes), f"{(nb_scannes/total_recues*100 if total_recues > 0 else 0):.1f}"],
        ["Nombre des moustiquaires non scannées", int(nb_non_scannes), f"{(nb_non_scannes/total_recues*100 if total_recues > 0 else 0):.1f}"],
        ["Nombre moustiquaire reçu", int(total_recues), "100"]
    ]
    create_table(doc, table_scan, ["Moustiquaire", "Effectif", "Fréquence (%)"])
    doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True
    
    

    if 'raison' in data.columns:
        doc.add_heading('Tableau : Raison de non Scannage', level=2)
        raison_series = data['raison'].str.split(', ').explode().value_counts()
        table_raison = [[v, c, f"{(c/total_resp*100):.1f}"] for v, c in raison_series.items()]
        create_table(doc, table_raison, ['Raison de non scannage', 'Effectif', '% de moustiquaires'])
        doc.add_paragraph('Source : Données issues du re-dénombrement 5% de la CDM-2026').italic = True
        
    # Sauvegarder en mémoire
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


################################################################################
# FONCTION POUR TÉLÉCHARGER LE RAPPORT
################################################################################

def download_automatic_report_button(data: pd.DataFrame, tables: dict):
    """Crée un bouton de téléchargement pour le rapport automatique"""
    
    st.markdown("---")
    st.markdown("### 📥 Téléchargement du rapport")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.info(
            "📄 Ce rapport contient une analyse complète selon la structure standard : "
            "caractéristiques, indicateurs de qualité, analyse de distribution, "
            "marquage, sensibilisation et recommandations."
        )
    
    with col2:
        if st.button("🔄 Générer le rapport", use_container_width=True):
            with st.spinner("Génération du rapport en cours..."):
                report_file = generate_automatic_report(data, tables)
                
                if report_file:
                    filename = f"Rapport_MILDA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    
                    st.download_button(
                        label="📥 Télécharger le rapport Word",
                        data=report_file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                    st.success("✅ Rapport généré avec succès !")
                else:
                    st.error("❌ Erreur lors de la génération du rapport")

################################################################################
# INTÉGRATION KOBO - FONCTIONS ROBUSTES
################################################################################

KOBO_SERVERS = {
    "KoBoToolbox Global — kf.kobotoolbox.org": "https://kf.kobotoolbox.org",
    "OCHA Humanitarian — kobo.humanitarianresponse.info": "https://kobo.humanitarianresponse.info",
    "Serveur personnalisé…": "custom",
}

def kobo_get_token(username, password, base_url):
    """Récupère le token API KoBo via Basic Auth"""
    url = f"{base_url}/token/?format=json"
    try:
        r = requests.get(url, auth=(username, password), timeout=20)
        r.raise_for_status()
        return r.json().get("token"), None
    except Exception as e:
        code = getattr(getattr(e, "response", None), "status_code", None)
        if code == 401:
            return None, "Identifiants KoBoToolbox invalides."
        return None, str(e)

def kobo_list_forms(token, base_url):
    """Liste tous les formulaires de type 'survey' disponibles"""
    url = f"{base_url}/api/v2/assets/?asset_type=survey&format=json"
    headers = {"Authorization": f"Token {token}"}
    try:
        r = requests.get(url, headers=headers, timeout=25)
        r.raise_for_status()
        data = r.json()
        forms = []
        for a in data.get("results", []):
            forms.append({
                "uid":         a.get("uid", ""),
                "name":        a.get("name", "Sans nom"),
                "submissions": a.get("deployment__submission_count", 0),
                "owner":       a.get("owner__username", ""),
                "modified":    str(a.get("date_modified", ""))[:10],
            })
        return forms, None
    except Exception as e:
        return None, str(e)

def kobo_download_data(token, uid, base_url):
    """Télécharge toutes les soumissions d'un formulaire (pagination automatique)"""
    headers = {"Authorization": f"Token {token}"}
    all_results = []
    next_url = f"{base_url}/api/v2/assets/{uid}/data/?format=json&limit=1000"
    try:
        while next_url:
            r = requests.get(next_url, headers=headers, timeout=60)
            r.raise_for_status()
            payload = r.json()
            batch = payload.get("results", [])
            all_results.extend(batch)
            next_url = payload.get("next")
        if not all_results:
            return pd.DataFrame(), None
        df = pd.json_normalize(all_results)
        df.columns = [c.split("/")[-1].lstrip("_") for c in df.columns]
        df = df.loc[:, ~df.columns.duplicated()]
        return df, None
    except Exception as e:
        return None, str(e)

def get_kobo_token(url, username, password):
    """Alias de compatibilité ascendante"""
    token, _ = kobo_get_token(username, password, url)
    return token

def process_raw_kobo_data(df):
    """Applique la logique de calcul des indicateurs sur les données brutes"""
    yes_no_cols = ['menage_servi', 'norme', 'menage_marque', 'information']
    for col in yes_no_cols:
        if col in df.columns:
            df[col] = df[col].apply(DataProcessor.normalize_yes_no)
    for col in ['nb_personnes', 'nb_milda_recues']:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')
    if 'nb_personnes' in df.columns:
        df['nb_milda_attendues'] = df['nb_personnes'].apply(DataProcessor.calculate_expected_milda)
    df['indic_servi']   = (df['menage_servi'].eq('Oui')).astype(int)  if 'menage_servi'  in df.columns else 0
    df['indic_correct'] = (df['norme'].eq('Oui')).astype(int)         if 'norme'         in df.columns else 0
    df['indic_marque']  = (df['menage_marque'].eq('Oui')).astype(int) if 'menage_marque' in df.columns else 0
    df['indic_info']    = (df['information'].eq('Oui')).astype(int)   if 'information'   in df.columns else 0
    if 'nb_milda_attendues' in df.columns and 'nb_milda_recues' in df.columns:
        df['ecart_distribution'] = df['nb_milda_recues'] - df['nb_milda_attendues']
    return df, {"total_rows": len(df)}
################################################################################
# APPLICATION PRINCIPALE
################################################################################

def main():
    """Fonction principale de l'application multisite et multi-phase"""
    
    # En-tête
    render_header()
    
    # --- CONSTANTE DES MAPPINGS GITHUB ---
    MAPPINGS_URLS = {
        "Phase 1": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_1.xlsx",
        "Phase 2": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_2.xlsx",
        "Phase 3": "https://github.com/dhp-byte/MILDA-Monitoring/raw/main/Choix_Phase_3.xlsx"
    }
    
    # Initialisation des variables de session indispensables
    for _k, _v in [('kobo_token', None), ('kobo_step', 0), ('kobo_forms', None),
                   ('kobo_selected', {}), ('data', None), ('tables', None), ('mappings', {})]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # Toutes les phases sont toujours actives (sélection automatique)
    selected_phases = list(MAPPINGS_URLS.keys())

    # --- 1. CONFIGURATION SIDEBAR ---
    with st.sidebar:
        # Connexion KoBo : déclenchement du wizard multi-étapes
        if st.button("🔗 Connexion KoBoToolbox", use_container_width=True, key="kobo_wizard_open"):
            st.session_state.update(kobo_step=1, kobo_forms=None, kobo_selected={})
            st.rerun()
        if st.session_state.get('kobo_token'):
            st.success("✅ KoBo connecté")
            if st.button("🔓 Déconnecter KoBo", key="kobo_logout"):
                for k in ['kobo_token', 'kobo_base_url', 'kobo_username', 'kobo_step',
                          'kobo_forms', 'kobo_selected']:
                    st.session_state.pop(k, None)
                st.rerun()

        st.divider()
        st.header("📂 Import Excel")
        uploaded_files = st.file_uploader(
            "Choisir un ou plusieurs fichiers Excel",
            type=['xlsx', 'xls'],
            accept_multiple_files=True
        )

    # --- 2. CHARGEMENT DYNAMIQUE DES MAPPINGS GITHUB EN FONCTION DES PHASES ---
    current_mappings_key = "-".join(sorted(selected_phases))
    if st.session_state.get('last_mappings_key') != current_mappings_key:
        with st.spinner("Chargement et fusion des référentiels cartographiques..."):
            combined_mappings = {}
            for phase in selected_phases:
                url = MAPPINGS_URLS.get(phase)
                if url:
                    try:
                        response = requests.get(url, timeout=15)
                        if response.status_code == 200:
                            df_choices = pd.read_excel(BytesIO(response.content), sheet_name='Choix', dtype=str)
                            df_choices.columns = df_choices.columns.str.strip()
                            df_choices = df_choices.dropna(subset=['list_name', 'value'])
                            for list_name in df_choices['list_name'].unique():
                                clean_list_key = str(list_name).strip()
                                subset = df_choices[df_choices['list_name'] == list_name]
                                phase_dict = dict(zip(subset['value'].str.strip(), subset['label'].str.strip()))
                                if clean_list_key in combined_mappings:
                                    combined_mappings[clean_list_key].update(phase_dict)
                                else:
                                    combined_mappings[clean_list_key] = phase_dict
                    except Exception as e:
                        st.sidebar.error(f"Erreur dictionnaire {phase} : {e}")
            st.session_state.mappings = combined_mappings
            st.session_state.last_mappings_key = current_mappings_key
            st.sidebar.success(f"✅ {len(combined_mappings)} listes fusionnées.")

    # --- 3. WIZARD KOBO MULTI-ÉTAPES ---
    kobo_step = st.session_state.get('kobo_step', 0)

    # ÉTAPE 1 : Authentification
    if kobo_step == 1:
        st.markdown("---")
        st.markdown("### 🔑 Connexion KoBoToolbox — Étape 1/3 : Authentification")
        col_srv, col_u, col_p = st.columns([2, 1, 1])
        with col_srv:
            srv_choice = st.selectbox("Serveur KoBo", list(KOBO_SERVERS.keys()), key="kobo_srv")
            base_url = KOBO_SERVERS[srv_choice]
            if base_url == "custom":
                base_url = st.text_input("URL du serveur personnalisé", placeholder="https://votre-kobo.org", key="kobo_cust")
        with col_u:
            kobo_u = st.text_input("Nom d'utilisateur", placeholder="utilisateur", key="kobo_u")
        with col_p:
            kobo_p = st.text_input("Mot de passe", type="password", placeholder="••••••••", key="kobo_p")

        st.info("🔒 Les identifiants sont stockés uniquement en session navigateur et effacés à la déconnexion.")
        ca, cb = st.columns([3, 1])
        with ca:
            if st.button("Se connecter →", use_container_width=True, key="kobo_connect"):
                if not kobo_u or not kobo_p:
                    st.error("Veuillez saisir votre nom d'utilisateur et votre mot de passe.")
                elif not base_url or base_url == "custom":
                    st.error("Veuillez saisir une URL de serveur valide.")
                else:
                    with st.spinner("Authentification en cours…"):
                        token, err = kobo_get_token(kobo_u, kobo_p, base_url)
                    if err:
                        st.error(f"❌ {err}")
                    else:
                        st.session_state.update(kobo_token=token, kobo_base_url=base_url,
                                                kobo_username=kobo_u, kobo_step=2)
                        st.rerun()
        with cb:
            if st.button("← Annuler", key="kobo_cancel1"):
                st.session_state['kobo_step'] = 0
                st.rerun()
        st.stop()

    # ÉTAPE 2 : Sélection des formulaires
    elif kobo_step == 2:
        token    = st.session_state.get('kobo_token', '')
        base_url = st.session_state.get('kobo_base_url', '')
        kuser    = st.session_state.get('kobo_username', '')

        st.markdown("---")
        st.markdown(f"### 📋 Connexion KoBoToolbox — Étape 2/3 : Sélection des formulaires")
        st.caption(f"Connecté en tant que **{kuser}** sur {base_url}")

        if 'kobo_forms' not in st.session_state or st.session_state.kobo_forms is None:
            with st.spinner("Récupération de la liste des formulaires…"):
                forms, err = kobo_list_forms(token, base_url)
            if err:
                st.error(f"❌ {err}")
                if st.button("← Retour", key="k2b_err"):
                    st.session_state['kobo_step'] = 1
                    st.rerun()
                st.stop()
            st.session_state['kobo_forms'] = forms or []

        forms = st.session_state.get('kobo_forms', [])
        if not forms:
            st.warning("Aucun formulaire trouvé. Vérifiez vos accès KoBo.")
            if st.button("← Retour", key="k2b_empty"):
                st.session_state['kobo_step'] = 1
                st.rerun()
            st.stop()

        search   = st.text_input("🔍 Filtrer les formulaires par nom", placeholder="Rechercher…", key="fsearch")
        filtered = [f for f in forms if search.lower() in f["name"].lower()] if search else forms
        st.caption(f"{len(filtered)} formulaire(s) affiché(s)")

        selected = st.session_state.get('kobo_selected', {})
        for form in filtered:
            uid  = form["uid"]
            name = form["name"]
            subs = form["submissions"]
            owner = form["owner"]
            mod  = form["modified"]
            cc, ci = st.columns([0.05, 0.95])
            with cc:
                checked = st.checkbox("", value=selected.get(uid, False), key=f"chk_{uid}")
                selected[uid] = checked
            with ci:
                color = "#10b981" if subs > 0 else "#64748B"
                st.markdown(
                    f"**📋 {name}** &nbsp;&nbsp;"
                    f"<span style='color:{color};font-size:.85rem;'>{subs:,} soumissions</span>"
                    f"<br><small>UID: `{uid}` · Propriétaire: {owner} · Modifié: {mod}</small>",
                    unsafe_allow_html=True
                )
        st.session_state['kobo_selected'] = selected

        n_sel = sum(1 for v in selected.values() if v)
        ca2, cb2, cc2 = st.columns([1, 2, 1])
        with ca2:
            if st.button("← Retour", key="k2_back"):
                st.session_state.pop('kobo_forms', None)
                st.session_state['kobo_step'] = 1
                st.rerun()
        with cb2:
            st.markdown(f"<div style='text-align:center;padding:.5rem;'><b>{n_sel}</b> formulaire(s) sélectionné(s)</div>", unsafe_allow_html=True)
        with cc2:
            if st.button(f"Suivant →", use_container_width=True, key="k2_load", disabled=(n_sel == 0)):
                st.session_state['kobo_step'] = 25
                st.rerun()
        st.stop()

    # ÉTAPE 2.5 : Attribution des phases aux formulaires
    elif kobo_step == 25:
        forms    = st.session_state.get('kobo_forms', [])
        selected = st.session_state.get('kobo_selected', {})
        sel_uids = [u for u, v in selected.items() if v]
        fnames   = {f["uid"]: f["name"] for f in forms}
        phases_dispo = list(MAPPINGS_URLS.keys())  # ["Phase 1", "Phase 2", "Phase 3"]

        st.markdown("---")
        st.markdown("### 🏷️ Connexion KoBoToolbox — Étape 2b/3 : Attribution des phases")
        st.caption("Indiquez à quelle phase appartient chaque formulaire sélectionné.")

        # Couleurs par phase pour un repère visuel
        phase_colors = {"Phase 1": "#3b82f6", "Phase 2": "#10b981", "Phase 3": "#f59e0b"}

        phase_assignments = st.session_state.get('kobo_phase_assignments', {})

        st.markdown("<br>", unsafe_allow_html=True)
        for uid in sel_uids:
            nm = fnames.get(uid, uid)
            col_name, col_sel = st.columns([2, 1])
            with col_name:
                st.markdown(f"**📋 {nm}**<br><small style='color:#64748b;'>{uid}</small>", unsafe_allow_html=True)
            with col_sel:
                # Tente de pré-détecter la phase depuis le nom du formulaire
                default_idx = 0
                for i, p in enumerate(phases_dispo):
                    if p.lower() in nm.lower() or f"phase{i+1}" in nm.lower().replace(" ", ""):
                        default_idx = i
                        break
                # Si déjà assigné, reprendre la valeur
                if uid in phase_assignments and phase_assignments[uid] in phases_dispo:
                    default_idx = phases_dispo.index(phase_assignments[uid])
                chosen = st.selectbox(
                    "Phase",
                    options=phases_dispo,
                    index=default_idx,
                    key=f"phase_sel_{uid}",
                    label_visibility="collapsed"
                )
                color = phase_colors.get(chosen, "#6b7280")
                st.markdown(
                    f"<div style='background:{color}22;border-left:3px solid {color};"
                    f"padding:2px 8px;border-radius:4px;font-size:.8rem;color:{color};font-weight:600;'>"
                    f"{chosen}</div>",
                    unsafe_allow_html=True
                )
                phase_assignments[uid] = chosen
            st.divider()

        st.session_state['kobo_phase_assignments'] = phase_assignments

        ca3, cb3 = st.columns([1, 3])
        with ca3:
            if st.button("← Retour", key="k25_back"):
                st.session_state['kobo_step'] = 2
                st.rerun()
        with cb3:
            if st.button("⬇️ Télécharger et traiter →", use_container_width=True, key="k25_next"):
                st.session_state['kobo_step'] = 3
                st.rerun()
        st.stop()

    # ÉTAPE 3 : Téléchargement et traitement
    elif kobo_step == 3:
        token    = st.session_state.get('kobo_token', '')
        base_url = st.session_state.get('kobo_base_url', '')
        forms    = st.session_state.get('kobo_forms', [])
        selected = st.session_state.get('kobo_selected', {})
        sel_uids = [u for u, v in selected.items() if v]
        fnames   = {f["uid"]: f["name"] for f in forms}
        phase_assignments = st.session_state.get('kobo_phase_assignments', {})

        st.markdown("---")
        st.markdown(f"### ⬇️ Connexion KoBoToolbox — Étape 3/3 : Téléchargement")
        st.caption(f"Téléchargement de {len(sel_uids)} formulaire(s) avec pagination complète…")

        pb   = st.progress(0)
        stat = st.empty()
        dfs_k = {}; logs = []

        for i, uid in enumerate(sel_uids):
            nm = fnames.get(uid, uid)
            pb.progress(i / len(sel_uids), text=f"Téléchargement : {nm}…")
            stat.info(f"⏳ En cours : {nm}…")
            df, err = kobo_download_data(token, uid, base_url)
            if err:
                logs.append(f"❌ **{nm}** — {err}")
            else:
                dfs_k[uid] = (nm, df if df is not None else pd.DataFrame())
                n_rec = len(df) if df is not None else 0
                n_col = len(df.columns) if df is not None else 0
                logs.append(f"✅ **{nm}** — {n_rec:,} enregistrements, {n_col} champs")

        pb.progress(1.0, text="Terminé !")
        stat.empty()
        for log in logs:
            st.markdown(log)

        if dfs_k:
            # ── Traitement : stockage SÉPARÉ par phase ────────────────────────
            phase_buckets = {}   # {"Phase 1": [df, df, …], "Phase 2": […], …}
            all_dfs       = []   # consolidation globale

            # Charger les mappings de chaque phase nécessaire
            phases_needed = set(phase_assignments.values())
            phase_mappings_cache = {}
            for phase in phases_needed:
                url = MAPPINGS_URLS.get(phase)
                if url:
                    try:
                        resp = requests.get(url, timeout=15)
                        if resp.status_code == 200:
                            df_ch = pd.read_excel(BytesIO(resp.content), sheet_name='Choix', dtype=str)
                            df_ch.columns = df_ch.columns.str.strip()
                            df_ch = df_ch.dropna(subset=['list_name', 'value'])
                            pm = {}
                            for ln in df_ch['list_name'].unique():
                                sub = df_ch[df_ch['list_name'] == ln]
                                pm[str(ln).strip()] = dict(zip(sub['value'].str.strip(), sub['label'].str.strip()))
                            phase_mappings_cache[phase] = pm
                            st.info(f"📖 Dictionnaire {phase} chargé ({len(pm)} listes)")
                    except Exception as e:
                        st.warning(f"⚠️ Impossible de charger le dictionnaire {phase} : {e}")
                        phase_mappings_cache[phase] = {}

            for uid, (fnm, df_raw) in dfs_k.items():
                detected_phase = phase_assignments.get(uid, selected_phases[0])
                phase_map = phase_mappings_cache.get(detected_phase, {})
                with st.spinner(f"Traitement de {fnm} ({detected_phase})…"):
                    data_proc, _ = process_milda_dataframe(df_raw, mappings_dict=phase_map)
                data_proc['phase_app'] = detected_phase
                phase_buckets.setdefault(detected_phase, []).append(data_proc)
                all_dfs.append(data_proc)

            # Stocker data_phase_1 / data_phase_2 / data_phase_3 individuellement
            phase_key_map = {
                "Phase 1": "data_phase_1",
                "Phase 2": "data_phase_2",
                "Phase 3": "data_phase_3",
            }
            for phase, key in phase_key_map.items():
                if phase in phase_buckets:
                    df_phase = pd.concat(phase_buckets[phase], ignore_index=True)
                    st.session_state[key]              = df_phase
                    st.session_state[key + "_tables"]  = generate_analysis_tables(df_phase)
                else:
                    # Conserver les données déjà en session si la phase n'a pas été rechargée
                    st.session_state.setdefault(key, None)
                    st.session_state.setdefault(key + "_tables", None)

            # Consolidation globale (toutes phases)
            consolidated = pd.concat(all_dfs, ignore_index=True)
            st.session_state['data_all']        = consolidated
            st.session_state['data_all_tables'] = generate_analysis_tables(consolidated)

            # Rétrocompatibilité : st.session_state.data = vue consolidée
            st.session_state.data   = consolidated
            st.session_state.tables = st.session_state['data_all_tables']

            st.session_state['kobo_step'] = 0

            # Résumé par phase
            for phase, dfs in phase_buckets.items():
                n = sum(len(d) for d in dfs)
                st.success(f"✅ {phase} — {n:,} enregistrements chargés")
            st.info(f"📦 Total consolidé : {len(consolidated):,} enregistrements")

            if st.button("Ouvrir le Dashboard →", use_container_width=True, key="k3_open"):
                st.rerun()
        else:
            st.error("Aucune donnée chargée. Vérifiez les permissions sur les formulaires sélectionnés.")
            if st.button("← Recommencer", key="k3_reset"):
                st.session_state['kobo_step'] = 0
                st.rerun()
        st.stop()

    # --- 5. LOGIQUE IMPORT EXCEL ---
    if uploaded_files and st.session_state.data is None:
        with st.spinner("🔄 Traitement et consolidation des fichiers Excel..."):
            phase_buckets_xl = {}
            all_dfs_xl = []
            for file in uploaded_files:
                df_temp, _ = load_and_process_data(file)
                if not df_temp.empty:
                    detected_phase = selected_phases[0]
                    for p in selected_phases:
                        if p.lower() in file.name.lower():
                            detected_phase = p
                            break
                    df_temp['phase_app'] = detected_phase
                    phase_buckets_xl.setdefault(detected_phase, []).append(df_temp)
                    all_dfs_xl.append(df_temp)

            if all_dfs_xl:
                phase_key_map = {
                    "Phase 1": "data_phase_1",
                    "Phase 2": "data_phase_2",
                    "Phase 3": "data_phase_3",
                }
                # Recharger les mappings par phase pour l'import Excel aussi
                phase_mappings_xl = {}
                for phase in phase_buckets_xl:
                    url = MAPPINGS_URLS.get(phase)
                    if url:
                        try:
                            resp = requests.get(url, timeout=15)
                            if resp.status_code == 200:
                                df_ch = pd.read_excel(BytesIO(resp.content), sheet_name='Choix', dtype=str)
                                df_ch.columns = df_ch.columns.str.strip()
                                df_ch = df_ch.dropna(subset=['list_name', 'value'])
                                pm = {}
                                for ln in df_ch['list_name'].unique():
                                    sub = df_ch[df_ch['list_name'] == ln]
                                    pm[str(ln).strip()] = dict(zip(sub['value'].str.strip(), sub['label'].str.strip()))
                                phase_mappings_xl[phase] = pm
                        except Exception:
                            phase_mappings_xl[phase] = {}

                for phase, key in phase_key_map.items():
                    if phase in phase_buckets_xl:
                        phase_map = phase_mappings_xl.get(phase, {})
                        # Re-traiter chaque df brut avec le bon mapping
                        dfs_reprocessed = []
                        for df_p_raw in phase_buckets_xl[phase]:
                            df_p_proc, _ = process_milda_dataframe(df_p_raw, mappings_dict=phase_map)
                            dfs_reprocessed.append(df_p_proc)
                        df_p = pd.concat(dfs_reprocessed, ignore_index=True)
                        st.session_state[key]             = df_p
                        st.session_state[key + "_tables"] = generate_analysis_tables(df_p)
                    else:
                        st.session_state.setdefault(key, None)
                        st.session_state.setdefault(key + "_tables", None)

                consolidated_data = pd.concat(all_dfs_xl, ignore_index=True)
                st.session_state['data_all']        = consolidated_data
                st.session_state['data_all_tables'] = generate_analysis_tables(consolidated_data)
                st.session_state.data   = consolidated_data
                st.session_state.tables = st.session_state['data_all_tables']
                st.rerun()

    # --- 6. AFFICHAGE DES ONGLETS (Si données présentes) ---

    # ── Déterminer quelles phases sont disponibles ────────────────────────────
    phase_key_map = {
        "Phase 1": "data_phase_1",
        "Phase 2": "data_phase_2",
        "Phase 3": "data_phase_3",
    }
    phases_chargees = [p for p, k in phase_key_map.items()
                       if st.session_state.get(k) is not None]
    any_data = bool(phases_chargees) or st.session_state.data is not None

    if any_data:
        # ── Bouton effacer ────────────────────────────────────────────────────
        if st.sidebar.button("🗑️ Effacer les données de l'application"):
            for k in ['data', 'tables', 'data_all', 'data_all_tables',
                      'data_phase_1', 'data_phase_1_tables',
                      'data_phase_2', 'data_phase_2_tables',
                      'data_phase_3', 'data_phase_3_tables']:
                st.session_state.pop(k, None)
            st.rerun()

        # ── Sélecteur de phase ────────────────────────────────────────────────
        st.markdown("---")
        options_vue = phases_chargees.copy()
        if len(phases_chargees) > 1:
            options_vue += ["Toutes les phases"]

        # Icônes et couleurs par option
        phase_icons = {
            "Phase 1": "🔵", "Phase 2": "🟢", "Phase 3": "🟠",
            "Toutes les phases": "🌐"
        }
        phase_colors = {
            "Phase 1": "#3b82f6", "Phase 2": "#10b981",
            "Phase 3": "#f59e0b", "Toutes les phases": "#6366f1"
        }

        # Affichage en colonnes-boutons pour choisir la vue
        st.markdown("#### 📊 Choisir la vue à afficher :")
        btn_cols = st.columns(len(options_vue))
        if 'vue_active' not in st.session_state or st.session_state.vue_active not in options_vue:
            st.session_state.vue_active = options_vue[0]

        for i, opt in enumerate(options_vue):
            with btn_cols[i]:
                color  = phase_colors.get(opt, "#6b7280")
                icon   = phase_icons.get(opt, "📊")
                is_sel = st.session_state.vue_active == opt
                border = f"3px solid {color}" if is_sel else "1px solid #e2e8f0"
                bg     = f"{color}18" if is_sel else "white"
                st.markdown(
                    f"<div style='border:{border};background:{bg};border-radius:10px;"
                    f"padding:.6rem;text-align:center;'>"
                    f"<span style='font-size:1.4rem;'>{icon}</span><br>"
                    f"<span style='font-size:.85rem;font-weight:{'700' if is_sel else '400'};"
                    f"color:{color if is_sel else '#374151'};'>{opt}</span></div>",
                    unsafe_allow_html=True
                )
                if st.button(f"Voir {opt}", key=f"vue_btn_{i}", use_container_width=True):
                    st.session_state.vue_active = opt
                    st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)
        vue = st.session_state.vue_active

        # ── Résolution data_active / tables_active selon la vue ───────────────
        if vue == "Toutes les phases":
            data_active   = st.session_state.get('data_all') or st.session_state.data
            tables_active = st.session_state.get('data_all_tables') or st.session_state.tables
        else:
            key = phase_key_map.get(vue)
            data_active   = st.session_state.get(key)
            tables_active = st.session_state.get(key + "_tables")
            if data_active is None:
                st.warning(f"Aucune donnée disponible pour {vue}.")
                st.stop()

        color_vue = phase_colors.get(vue, "#6366f1")
        icon_vue  = phase_icons.get(vue, "📊")
        n_rec     = len(data_active) if data_active is not None else 0
        st.markdown(
            f"<div style='background:{color_vue}12;border-left:4px solid {color_vue};"
            f"border-radius:8px;padding:.6rem 1rem;margin-bottom:1rem;'>"
            f"<b style='color:{color_vue};'>{icon_vue} {vue}</b>"
            f" &nbsp;·&nbsp; <span style='color:#374151;'>{n_rec:,} enregistrements</span></div>",
            unsafe_allow_html=True
        )

        # ── Menu à onglets ────────────────────────────────────────────────────
        tabs = st.tabs([
            "🏠 Dashboard", "🔍 Analyse", "🗺️ Cartographie",
            "🏃 Suivi Agents", "🛡️ Qualité", "📊 Statistiques",
            "📥 Export", "📥 Rapport Final"
        ])

        with tabs[0]: page_dashboard(data_active, tables_active)
        with tabs[1]: page_analysis(data_active, tables_active)
        with tabs[2]: page_maps(data_active)
        with tabs[3]: page_agent_tracking(data_active)
        with tabs[4]: page_data_quality(data_active)
        with tabs[5]: page_statistics(data_active)
        with tabs[6]: page_export(data_active, tables_active)
        with tabs[7]:
            st.markdown("## 📊 Rapport de Synthèse Automatique")
            download_automatic_report_button(data_active, tables_active)

    else:
        # Écran d'accueil informatif
        st.info("👆 Connectez-vous à KoBo ou importez un ou plusieurs fichiers Excel (Sidebar) pour générer les analyses.")

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 📋 Structure attendue")
            st.code("province, district, village\ndate_enquete\nmenage_servi (Oui/Non)\nnb_personnes\nnb_milda_recues", language="text")
        with col2:
            st.markdown("### ⚙️ Paramètres système détectés")
            st.write(f"Version DOCX : {'✅ OK' if DOCX_AVAILABLE else '❌ Absent'}")
            st.write(f"Version Stats : {'✅ OK' if STATS_AVAILABLE else '❌ Absent'}")

    # Pied de page
    st.markdown("---")
    st.caption(f"🦟 Hub National Multi-Phases MILDA | Extraction synchronisée le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")

if __name__ == "__main__":
    main()
