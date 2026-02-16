################################################################################
# APPLICATION STREAMLIT - Monitorage externe MILDA
# Traduction Python du script R original
################################################################################

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import io
import zipfile
import re
import math
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration de la page [cite: 1, 2]
st.set_page_config(
    page_title="Rapport Monitorage MILDA",
    page_icon="🦟",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================
# FONCTIONS UTILITAIRES
# =========================

def yn(x):
    """Normalise les réponses Oui/Non"""
    if pd.isna(x):
        return np.nan
    x = str(x).strip().lower()
    if x in ['oui', 'yes', 'y', '1', 'true']:
        return 'Oui'
    elif x in ['non', 'no', 'n', '0', 'false']:
        return 'Non'
    else:
        return x #[cite: 2, 3]

def milda_attendues(n_personnes):
    """Calcule le nombre de MILDA attendues (1 MILDA pour 2 personnes)"""
    try:
        n = float(n_personnes)
        if pd.isna(n):
            return np.nan
        return math.ceil(n / 2) #[cite: 3]
    except:
        return np.nan

def safe_filename(x):
    """Crée un nom de fichier sécurisé"""
    x = str(x)
    x = re.sub(r'[^A-Za-z0-9]+', '_', x)
    x = re.sub(r'^_+|_+$', '', x)
    return x #[cite: 4]

def create_bar_chart(df, x_col, y_cols, title, subtitle, colors=None):
    """Crée un graphique à barres avec Plotly"""
    fig = go.Figure()
    if colors is None:
        colors = px.colors.qualitative.Set2
    for i, col in enumerate(y_cols):
        fig.add_trace(go.Bar(
            y=df[x_col],
            x=df[col],
            name=col,
            orientation='h',
            marker_color=colors[i % len(colors)],
            text=df[col].apply(lambda x: f'{x:.1f}%'),
            textposition='outside'
        )) #[cite: 4, 5]
    fig.update_layout(
        title={
            'text': f"<b>{title}</b><br><sub>{subtitle}</sub>",
            'x': 0.5,
            'xanchor': 'center'
        },
        xaxis_title="Pourcentage",
        yaxis_title="",
        barmode='group',
        height=max(400, len(df) * 50),
        showlegend=True,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5),
        template="plotly_white"
    ) #[cite: 6, 7]
    return fig

def add_table_to_doc(doc, df):
    """Ajoute un DataFrame comme tableau dans le document Word"""
    if df.empty:
        doc.add_paragraph("Aucune donnée disponible")
        return
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        cell.paragraphs[0].runs[0].font.bold = True
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.rows[i + 1].cells[j].text = str(value)
    doc.add_paragraph() #[cite: 13, 14]

def create_docx_report(data, tables, graphs_bytes):
    """Génère le rapport Word (DOCX)"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    title = doc.add_heading('RAPPORT DE MONITORAGE EXTERNE – MILDA', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Analyse suivant le plan d\'analyse (indicateurs de qualité du dénombrement-distribution)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(datetime.now().strftime('%d/%m/%Y'))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break() #[cite: 7, 8]
    
    doc.add_heading('1. Indicateurs de qualité du dénombrement-distribution', level=1)
    doc.add_paragraph('Indicateurs suivis : % ménages servis, % ménages correctement servis, % ménages coupon, % marquage, % information utilisation correcte.') #[cite: 9]
    
    doc.add_heading('1.1 Pourcentage des ménages servis et correctement servis par Province', level=2)
    if 'G1' in graphs_bytes:
        doc.add_picture(io.BytesIO(graphs_bytes['G1']), width=Inches(6.5)) #[cite: 9]
    
    doc.add_heading('1.2 Tableau résumé des indicateurs par Province', level=2)
    add_table_to_doc(doc, tables['T0_resume_prov']) #[cite: 10]
    
    doc.add_heading('1.3 Tableaux par District sanitaire (DS) – par Province', level=2)
    provinces = sorted(data['province'].dropna().unique())
    for province in provinces:
        doc.add_heading(f'Province : {province}', level=3)
        doc.add_heading('A) Ménages servis et correctement servis (par DS)', level=4)
        t1p = tables['T1_servis_ds'][tables['T1_servis_ds']['province'] == province].drop('province', axis=1)
        add_table_to_doc(doc, t1p)
        doc.add_heading('B) Marquage des ménages (par DS)', level=4)
        t2p = tables['T2_marque_ds'][tables['T2_marque_ds']['province'] == province].drop('province', axis=1)
        add_table_to_doc(doc, t2p)
        doc.add_heading('C) Information sur l\'utilisation correcte des MILDA (par DS)', level=4)
        t3p = tables['T3_info_ds'][tables['T3_info_ds']['province'] == province].drop('province', axis=1)
        add_table_to_doc(doc, t3p) #[cite: 10, 11, 12]
    
    doc.add_heading('1.4 Pourcentage de ménages avec marquage par Province', level=2)
    if 'G2' in graphs_bytes:
        doc.add_picture(io.BytesIO(graphs_bytes['G2']), width=Inches(6.5)) #[cite: 12]
    
    doc.add_heading('1.5 Pourcentage de ménages ayant reçu l\'information (utilisation correcte) par Province', level=2)
    if 'G3' in graphs_bytes:
        doc.add_picture(io.BytesIO(graphs_bytes['G3']), width=Inches(6.5)) #[cite: 12, 13]
    return doc

# =========================
# INTERFACE STREAMLIT
# =========================

st.title("Rapport de Monitorage Externe – MILDA")
st.markdown("### Analyse des indicateurs de qualité du dénombrement-distribution")

with st.sidebar:
    st.header("Configuration")
    uploaded_file = st.file_uploader("Importer le fichier Excel", type=['xlsx', 'xls'])
    sheet_name = st.text_input("Nom de la feuille Excel", value="MONITORAGE EXTERNE DU DENOMB...")
    st.markdown("---")
    st.markdown("**Indicateurs suivis:**")
    st.markdown("- % ménages servis\n- % ménages correctement servis\n- % ménages coupon\n- % ménages marquage\n- % ménages info utilisation correcte") #[cite: 15, 16]

if uploaded_file is None:
    st.info("Veuillez importer un fichier Excel pour commencer l'analyse")
    st.stop()

# =========================
# TRAITEMENT DES DONNÉES
# =========================

try:
    with st.spinner("Chargement et traitement des données..."):
        try:
            data_raw = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        except:
            data_raw = pd.read_excel(uploaded_file, sheet_name=0) #[cite: 17]
        
        data = data_raw.rename(columns={
            'Province': 'province',
            'District sanitaire de :': 'district',
            'Est-ce que le ménage a-t-il été servis en MILDA lors de la campagne de distribution de masse ?': 'menage_servi',
            'Nombre des personnes qui habitent dans le ménage': 'nb_personnes',
            'Combien de MILDA avez-vous reçues ?': 'nb_milda_recues',
            'verif_cle': 'verif_cle',
            'Est-ce que le ménage a  été marqué comme un ménage ayant reçu de MILDA?': 'menage_marque',
            'Avez-vous été sensibilisé sur l\'utilisation correcte du MILDA par les relais communautaires ?': 'sensibilise'
        }) #[cite: 18, 19, 20]
        
        for col in ['menage_servi', 'verif_cle', 'menage_marque', 'sensibilise']:
            data[col] = data[col].apply(yn) #[cite: 20]
        
        data['nb_personnes'] = pd.to_numeric(data['nb_personnes'], errors='coerce')
        data['indic_servi'] = data['menage_servi'].apply(lambda x: 1 if x == 'Oui' else (0 if x == 'Non' else np.nan))
        data['indic_correct'] = data.apply(lambda row: 1 if (row['menage_servi'] == 'Oui' and row['verif_cle'] == 'Oui') else (0 if row['menage_servi'] == 'Oui' else np.nan), axis=1)
        data['indic_marque'] = data.apply(lambda row: 1 if (row['menage_servi'] == 'Oui' and row['menage_marque'] == 'Oui') else (0 if row['menage_servi'] == 'Oui' else np.nan), axis=1)
        data['indic_info'] = data['sensibilise'].apply(lambda x: 1 if x == 'Oui' else (0 if x == 'Non' else np.nan)) #[cite: 21, 22, 23, 24]

        # T0: Résumé par province
        T0_resume_prov = data.groupby('province').agg(
            menages_redenombres=('province', 'count'),
            pct_servis=('indic_servi', lambda x: round(100 * (x == 1).mean(), 1)),
            pct_correct=('indic_correct', lambda x: round(100 * (x == 1).mean(), 1)),
            pct_marque=('indic_marque', lambda x: round(100 * (x == 1).mean(), 1)),
            pct_info=('indic_info', lambda x: round(100 * (x == 1).mean(), 1))
        ).reset_index() #[cite: 24, 25]

        # Tableaux détaillés (T1, T2, T3)
        T1_servis_ds = data.groupby(['province', 'district']).agg(
            nb_redenombres=('district', 'count'),
            pct_servis=('indic_servi', lambda x: round(100 * (x == 1).mean(), 1)),
            pct_correct=('indic_correct', lambda x: round(100 * (x == 1).mean(), 1))
        ).reset_index()
        
        T2_marque_ds = data.groupby(['province', 'district']).agg(
            pct_marque=('indic_marque', lambda x: round(100 * (x == 1).mean(), 1))
        ).reset_index()

        T3_info_ds = data.groupby(['province', 'district']).agg(
            pct_info=('indic_info', lambda x: round(100 * (x == 1).mean(), 1))
        ).reset_index()
        
        tables = {'T0_resume_prov': T0_resume_prov, 'T1_servis_ds': T1_servis_ds, 'T2_marque_ds': T2_marque_ds, 'T3_info_ds': T3_info_ds} #[cite: 27, 28, 29, 30, 31]

        # Graphiques
        tab_servi_prov = T0_resume_prov.sort_values('pct_servis')
        G1 = create_bar_chart(tab_servi_prov, 'province', ['pct_servis', 'pct_correct'], 'Ménages servis et correctement servis', 'Par Province')
        
        G2 = go.Figure(go.Bar(y=T0_resume_prov['province'], x=T0_resume_prov['pct_marque'], orientation='h'))
        G3 = go.Figure(go.Bar(y=T0_resume_prov['province'], x=T0_resume_prov['pct_info'], orientation='h'))
        graphs = {'G1': G1, 'G2': G2, 'G3': G3} #[cite: 32, 33, 34, 35, 36, 37]

    # Affichage Streamlit
    tab1, tab2, tab3, tab4 = st.tabs(["Vue d'ensemble", "Graphiques", "Tableaux", "Téléchargements"]) #[cite: 38]
    
    with tab1:
        st.subheader("Résumé par Province")
        st.dataframe(T0_resume_prov, use_container_width=True) #[cite: 40]
        
    with tab2:
        st.plotly_chart(G1, use_container_width=True)
        st.plotly_chart(G2, use_container_width=True)
        st.plotly_chart(G3, use_container_width=True) #[cite: 40]

    with tab3:
        prov = st.selectbox("Filtrer par Province", ['Toutes'] + sorted(data['province'].unique().tolist())) #[cite: 41]
        display_t1 = T1_servis_ds if prov == 'Toutes' else T1_servis_ds[T1_servis_ds['province'] == prov]
        st.dataframe(display_t1, use_container_width=True) #[cite: 42, 43]

    with tab4:
        if st.button("Générer le rapport Word complet"):
            graphs_bytes = {name: fig.to_image(format='png') for name, fig in graphs.items()}
            doc = create_docx_report(data, tables, graphs_bytes)
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            st.download_button("Télécharger le rapport Word", docx_buffer.getvalue(), "Rapport_MILDA.docx") #[cite: 59, 60, 61, 62]

    # FOOTER (Section corrigée)
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray; padding: 20px;'>
        <p>Rapport de Monitorage Externe MILDA - Version Python/Streamlit</p>
        <p>Généré le {}</p>
    </div>
    """.format(datetime.now().strftime('%d/%m/%Y à %H:%M')), unsafe_allow_html=True) #[cite: 63]

except Exception as e:
    st.error(f"Erreur lors du traitement : {str(e)}")
    st.exception(e) #[cite: 63]
